# -*- coding: utf-8 -*-
"""محرك حقن البيانات في القالب"""

from io import BytesIO
from docx import Document
from .schedule import ROUTINE_ACTIVITIES, DOMAIN_MAPPING


def _safe_replace(paragraph, old, new):
    if old not in paragraph.text:
        return False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    full = paragraph.text.replace(old, new)
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = full
    return True


def build_daily_planner_bytes(
    day: str,
    template_bytes: bytes,
    schedule: dict,
    lessons_db: dict,
    week_num: str = "",
    date_str: str = "",
) -> tuple:
    """
    بناء كراس يوم واحد
    Returns: (bytes, sessions_info, warnings)
    """
    plan = schedule.get(day, [])
    if not plan:
        return None, [], [f"اليوم '{day}' غير موجود"]

    doc = Document(BytesIO(template_bytes))

    replacements = {
        "{{اليوم}}":    day,
        "{{التاريخ}}":  date_str,
        "{{الأسبوع}}":  week_num,
    }

    sessions_info = []
    warnings = []

    for i, session in enumerate(plan, 1):
        act  = session["النشاط"]
        dur  = session["المدة"]
        per  = session.get("الفترة", "")

        # مفاتيح القالب
        k = {
            'نشاط':  f"{{{{نشاط_{i}}}}}",
            'موضوع': f"{{{{موضوع_{i}}}}}",
            'كفاءة': f"{{{{كفاءة_{i}}}}}",
            'مدة':   f"{{{{مدة_{i}}}}}",
            'ميدان': f"{{{{ميدان_{i}}}}}",
        }

        replacements[k['مدة']]  = dur
        replacements[k['نشاط']] = act

        # المجال التعليمي
        domain = DOMAIN_MAPPING.get(act, "—")

        info = {
            "رقم": i,
            "النشاط": act,
            "المدة": dur,
            "الفترة": per,
            "المجال": domain,
            "نوع": "روتيني",
            "الموضوع": "—",
            "الكفاءة": "—",
        }

        if act in ROUTINE_ACTIVITIES:
            replacements[k['موضوع']] = "—"
            replacements[k['كفاءة']] = "—"
            replacements[k['ميدان']] = "—"

        elif act in lessons_db and lessons_db[act]:
            lesson = lessons_db[act].pop(0)
            topic = lesson.get('موضوع', '—')
            indic = lesson.get('كفاءة', '—')

            replacements[k['موضوع']] = topic
            replacements[k['كفاءة']] = indic
            replacements[k['ميدان']] = domain

            info["نوع"]     = "تعليمي"
            info["الموضوع"] = topic
            info["الكفاءة"] = indic

        else:
            replacements[k['موضوع']] = "⚠ لا توجد مذكرة"
            replacements[k['كفاءة']] = "⚠ لا توجد مذكرة"
            replacements[k['ميدان']] = domain
            info["نوع"] = "ناقص"
            warnings.append(act)

        sessions_info.append(info)

    # الحقن في الجداول
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in para.text:
                            _safe_replace(para, key, str(val))

    # الحقن في الفقرات
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                _safe_replace(para, key, str(val))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), sessions_info, warnings
