# -*- coding: utf-8 -*-
"""محرك استخراج الدروس من ملفات المذكرات"""

import re
from io import BytesIO
from docx import Document
from .schedule import NAME_MAPPING


# ═══════════════════════════════════════
#  تنظيف النصوص
# ═══════════════════════════════════════

def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'ـ+', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[\.\s]+$', '', text)
    return text.strip()


def normalize_name(raw: str) -> str:
    """
    توحيد اسم المادة ليطابق التوقيت
    البحث يتم بثلاث طرق: مطابقة تامة → جزئية → إرجاع كما هو
    """
    cleaned = clean_text(raw)
    if not cleaned:
        return cleaned

    # 1) مطابقة تامة
    if cleaned in NAME_MAPPING:
        return NAME_MAPPING[cleaned]

    # 2) مطابقة بدون "ال" التعريف
    without_al = re.sub(r'^ال', '', cleaned)
    if without_al in NAME_MAPPING:
        return NAME_MAPPING[without_al]

    # 3) بحث جزئي (المادة تحتوي على المفتاح أو العكس)
    cleaned_lower = cleaned
    for key, val in NAME_MAPPING.items():
        if key in cleaned_lower or cleaned_lower in key:
            return val

    # 4) بحث بالكلمات الجذرية
    root_map = {
        "علم":      "ت علمية وتكنولوجية",
        "تكنولوج":  "ت علمية وتكنولوجية",
        "إسلام":    "ت إسلامية",
        "مدن":      "ت مدنية",
        "بدن":      "ت بدنية",
        "إيقاع":    "ت إيقاعية",
        "تشكيل":    "تربية تشكيلية",
        "رياض":     "رياضيات",
        "قراءة":    "مبادئ القراءة",
        "تخطيط":    "تخطيط",
        "كتابة":    "تخطيط",
        "شفوي":     "تعبير شفوي",
        "تعبير":    "تعبير شفوي",
        "مسرح":     "مسرح وعرائس",
        "موسيق":    "موسيقى وإنشاد",
        "إنشاد":    "موسيقى وإنشاد",
        "رسم":      "تربية تشكيلية",
        "أشغال":    "تربية تشكيلية",
    }
    for root, target in root_map.items():
        if root in cleaned:
            return target

    return cleaned


# ═══════════════════════════════════════
#  أنماط Regex مرنة
# ═══════════════════════════════════════

RE_ACT = re.compile(
    r'^(?:النشاط|المادة|مجال\s*التعل[ـم]*|الميدان)\s*[:/\-]\s*(.*)',
)
RE_TOP = re.compile(
    r'^(?:الموضوع|الوحدة|عنوان\s*الدرس|المحتوى)\s*[:/\-]\s*(.*)',
)
RE_IND = re.compile(
    r'^(?:مؤشر\s*الكفا[ـءئ]*ة|الكفاءة\s*المستهدفة'
    r'|مؤشرات?\s*الكفاءة)\s*[:/\-]\s*(.*)',
)


# ═══════════════════════════════════════
#  محرك الاستخراج
# ═══════════════════════════════════════

def _extract_paragraphs(doc) -> dict:
    lessons = {}
    cur_act = None
    cur_les = {}

    def _save():
        nonlocal cur_act, cur_les
        if cur_act and cur_les.get('موضوع'):
            name = normalize_name(cur_act)
            lessons.setdefault(name, []).append(cur_les.copy())

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        m = RE_ACT.search(text)
        if m:
            _save()
            cur_act = m.group(1).strip()
            cur_les = {}
            continue

        m = RE_TOP.search(text)
        if m:
            cur_les['موضوع'] = clean_text(m.group(1))
            continue

        m = RE_IND.search(text)
        if m:
            cur_les['كفاءة'] = clean_text(m.group(1))
            continue

    _save()
    return lessons


def _extract_tables(doc, existing: dict = None) -> dict:
    if existing is None:
        existing = {}

    for table in doc.tables:
        cur_act = None
        cur_les = {}

        for row in table.rows:
            for cell in row.cells:
                text = clean_text(cell.text)
                if not text:
                    continue

                m = RE_ACT.search(text)
                if m:
                    if cur_act and cur_les.get('موضوع'):
                        name = normalize_name(cur_act)
                        existing.setdefault(name, []).append(
                            cur_les.copy()
                        )
                    cur_act = m.group(1).strip()
                    cur_les = {}
                    continue

                m = RE_TOP.search(text)
                if m:
                    cur_les['موضوع'] = clean_text(m.group(1))
                    continue

                m = RE_IND.search(text)
                if m:
                    cur_les['كفاءة'] = clean_text(m.group(1))
                    continue

        if cur_act and cur_les.get('موضوع'):
            name = normalize_name(cur_act)
            existing.setdefault(name, []).append(cur_les.copy())

    return existing


def extract_all_lessons(file_bytes: bytes) -> dict:
    """
    استخراج كل الدروس من ملف المذكرات
    المدخل: bytes من ملف docx
    المخرج: {'اسم_المادة': [{'موضوع': ..., 'كفاءة': ...}, ...]}
    """
    doc = Document(BytesIO(file_bytes))
    lessons = _extract_paragraphs(doc)
    lessons = _extract_tables(doc, lessons)
    return lessons
