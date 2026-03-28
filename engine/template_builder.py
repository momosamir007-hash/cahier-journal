# -*- coding: utf-8 -*-
"""بناء قالب الكراس اليومي"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn('w:bidi'), {}))


def _cell(cell, text, bold=False, size=10, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Sakkal Majalla"
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rPr.append(rPr.makeelement(
        qn('w:rFonts'), {qn('w:cs'): 'Sakkal Majalla'}
    ))


def _shade(row, hex_color):
    for c in row.cells:
        tcPr = c._tc.get_or_add_tcPr()
        tcPr.append(tcPr.makeelement(qn('w:shd'), {
            qn('w:fill'): hex_color, qn('w:val'): 'clear'
        }))


def _period_table(doc, title, start, count):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl(h)
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 51, 102)

    headers = [
        'مؤشرات الكفاءة', 'عنوان الدرس',
        'الميدان', 'النشاط', 'المدة'
    ]

    tbl = doc.add_table(rows=1 + count, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Cm(5.5), Cm(5), Cm(3.5), Cm(3.5), Cm(2)]
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    hdr = tbl.rows[0]
    _shade(hdr, "1F4E79")
    for i, txt in enumerate(headers):
        _cell(hdr.cells[i], txt, bold=True, size=11,
              color=RGBColor(255, 255, 255))

    for j in range(count):
        n = start + j
        dr = tbl.rows[1 + j]
        if j % 2 == 0:
            _shade(dr, "EDF2F9")
        placeholders = [
            f'{{{{كفاءة_{n}}}}}',
            f'{{{{موضوع_{n}}}}}',
            f'{{{{ميدان_{n}}}}}',
            f'{{{{نشاط_{n}}}}}',
            f'{{{{مدة_{n}}}}}',
        ]
        for i, ph in enumerate(placeholders):
            _cell(dr.cells[i], ph, size=9)


def create_template_bytes() -> bytes:
    doc = Document()

    for sec in doc.sections:
        sectPr = sec._sectPr
        sectPr.append(sectPr.makeelement(qn('w:bidi'), {}))

    for txt, sz, b in [
        ('الجمهورية الجزائرية الديمقراطية الشعبية', 12, True),
        ('وزارة التربية الوطنية', 11, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _rtl(p)
        r = p.add_run(txt)
        r.bold = b
        r.font.size = Pt(sz)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl(tp)
    tr = tp.add_run('الكراس اليومي')
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0, 51, 102)

    info = doc.add_table(rows=1, cols=3)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell(info.rows[0].cells[2], 'اليوم : {{اليوم}}', True, 12)
    _cell(info.rows[0].cells[1], 'التاريخ : {{التاريخ}}', size=11)
    _cell(info.rows[0].cells[0], 'الأسبوع : {{الأسبوع}}', size=11)

    doc.add_paragraph('')
    _period_table(doc, '☀ الفترة الصباحية', 1, 7)
    doc.add_paragraph('')
    _period_table(doc, '🌙 الفترة المسائية', 8, 5)
    doc.add_paragraph('')

    np = doc.add_paragraph()
    np.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _rtl(np)
    nr = np.add_run(
        'ملاحظات : '
        '........................................................'
        '........................................................'
    )
    nr.font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
