# -*- coding: utf-8 -*-
"""
╔══════════════════════════════════════════════════╗
║ أتمتة الكراس اليومي - قسم التحضيري              ║
║ الإصدار: 2.0 (مُحسَّن)                          ║
╚══════════════════════════════════════════════════╝
الاستخدام: python daily_planner.py
المتطلبات: pip install python-docx
"""

import re
import os
import copy
from docx import Document

# ═══════════════════════════════════════════════════════
# القسم 1: الثوابت
# ═══════════════════════════════════════════════════════

ROUTINE_ACTIVITIES = [
    "الاستقبال",
    "الاستراحة",
    "تهيئة الخروج",
    "نهاية الخروج",
]

# قاموس توحيد الأسماء بين المذكرات والتوقيت
NAME_MAPPING = {
    # ── المذكرات ──────────── → ── التوقيت ──
    "تربية علمية": "ت علمية",
    "تربية علمية وتكنولوجية": "ت علمية",
    "تربية إسلامية": "ت إسلامية",
    "تربية مدنية": "ت مدنية",
    "تربية بدنية": "ت بدنية",
    "تربية بدنية ورياضية": "ت بدنية",
    "تربية تكنولوجية": "ت تكنولوجية",
    "تربية إيقاعية": "ت إيقاعية",
    "تربية موسيقية": "موسيقى وإنشاد",
    "قراءة": "مبادئ القراءة",
    "مبادئ في القراءة": "مبادئ القراءة",
    "الرياضيات": "رياضيات",
    "كتابة": "تخطيط",
    "خط": "تخطيط",
    "مسرح": "مسرح وعرائس",
    "رسم": "رسم وأشغال",
    "أشغال يدوية": "رسم وأشغال",
    "موسيقى": "موسيقى وإنشاد",
    "إنشاد": "موسيقى وإنشاد",
    "تعبير": "تعبير شفوي",
}

# ═══════════════════════════════════════════════════════
# القسم 2: التوقيت الأسبوعي
# ═══════════════════════════════════════════════════════

weekly_schedule = {
    "الأحد": [
        {"النشاط": "الاستقبال", "المدة": "15 د"},
        {"النشاط": "تعبير شفوي", "المدة": "30 د"},
        {"النشاط": "مبادئ القراءة", "المدة": "30 د"},
        {"النشاط": "رياضيات", "المدة": "30 د"},
        {"النشاط": "ت علمية", "المدة": "30 د"},
        {"النشاط": "ت إسلامية", "المدة": "30 د"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د"},
        {"النشاط": "الاستقبال", "المدة": "15 د"},
        {"النشاط": "مسرح وعرائس", "المدة": "30 د"},
        {"النشاط": "رسم وأشغال", "المدة": "30 د"},
        {"النشاط": "ت بدنية", "المدة": "30 د"},
        {"النشاط": "نهاية الخروج", "المدة": "15 د"},
    ],
    "الإثنين": [
        {"النشاط": "الاستقبال", "المدة": "15 د"},
        {"النشاط": "رياضيات", "المدة": "30 د"},
        {"النشاط": "تعبير شفوي", "المدة": "30 د"},
        {"النشاط": "تخطيط", "المدة": "30 د"},
        {"النشاط": "ت علمية", "المدة": "30 د"},
        {"النشاط": "ت مدنية", "المدة": "30 د"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د"},
        {"النشاط": "الاستقبال", "المدة": "15 د"},
        {"النشاط": "مسرح وعرائس", "المدة": "30 د"},
        {"النشاط": "رسم وأشغال", "المدة": "30 د"},
        {"النشاط": "ت بدنية", "المدة": "30 د"},
        {"النشاط": "نهاية الخروج", "المدة": "15 د"},
    ],
    "الثلاثاء": [
        {"النشاط": "الاستقبال", "المدة": "15 د"},
        {"النشاط": "تعبير شفوي", "المدة": "30 د"},
        {"النشاط": "مبادئ القراءة", "المدة": "30 د"},
        {"النشاط": "رياضيات", "المدة": "30 د"},
        {"النشاط": "ت إسلامية", "المدة": "30 د"},
        {"النشاط": "ت بدنية", "المدة": "30 د"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د"},
    ],
    "الأربعاء": [
        {"النشاط": "الاستقبال", "المدة": "15 د"},
        {"النشاط": "رياضيات", "المدة": "30 د"},
        {"النشاط": "مبادئ القراءة", "المدة": "30 د"},
        {"النشاط": "تخطيط", "المدة": "30 د"},
        {"النشاط": "ت علمية", "المدة": "30 د"},
        {"النشاط": "ت مدنية", "المدة": "30 د"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د"},
        {"النشاط": "الاستقبال", "المدة": "15 د"},
        {"النشاط": "ت إيقاعية", "المدة": "30 د"},
        {"النشاط": "موسيقى وإنشاد", "المدة": "30 د"},
        {"النشاط": "ت بدنية", "المدة": "30 د"},
        {"النشاط": "نهاية الخروج", "المدة": "15 د"},
    ],
    "الخميس": [
        {"النشاط": "الاستقبال", "المدة": "15 د"},
        {"النشاط": "مبادئ القراءة", "المدة": "30 د"},
        {"النشاط": "رياضيات", "المدة": "30 د"},
        {"النشاط": "ت علمية", "المدة": "30 د"},
        {"النشاط": "ت إيقاعية", "المدة": "30 د"},
        {"النشاط": "موسيقى وإنشاد", "المدة": "30 د"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د"},
    ],
}

# ═══════════════════════════════════════════════════════
# القسم 3: دوال المعالجة النصية
# ═══════════════════════════════════════════════════════


def clean_text(text):
    """تنظيف النص العربي"""
    if not text:
        return ""
    # إزالة التطويل
    text = re.sub(r'ـ+', '', text)
    # إزالة مسافات متعددة
    text = re.sub(r'\s+', ' ', text)
    # إزالة نقطة أو فراغ في النهاية
    text = re.sub(r'[\.\s]+$', '', text)
    return text.strip()


def normalize_name(raw_name):
    """توحيد اسم المادة"""
    cleaned = clean_text(raw_name)
    # بحث مباشر
    if cleaned in NAME_MAPPING:
        return NAME_MAPPING[cleaned]
    # بحث بدون تشكيل أو مسافات إضافية
    cleaned_compact = re.sub(r'\s+', ' ', cleaned)
    for key, value in NAME_MAPPING.items():
        # مقارنة مرنة
        if key in cleaned_compact or cleaned_compact in key:
            return value
    return cleaned


# ═══════════════════════════════════════════════════════
# القسم 4: محرك استخراج الدروس
# ═══════════════════════════════════════════════════════

# أنماط البحث المرنة (تتعامل مع التطويل والتشكيل)
RE_ACTIVITY = re.compile(
    r'^(?:النشاط|المادة|مجال\s*التعل[ـم]*)\s*[:/\-]\s*(.*)',
)
RE_TOPIC = re.compile(
    r'^(?:الموضوع|الوحدة|عنوان\s*الدرس)\s*[:/\-]\s*(.*)',
)
RE_INDICATOR = re.compile(
    r'^(?:مؤشر\s*الكفا[ـءئ]*ة|الكفاءة\s*المستهدفة)\s*[:/\-]\s*(.*)',
)


def extract_from_paragraphs(doc):
    """استخراج الدروس من فقرات المستند"""
    lessons = {}
    current_activity = None
    current_lesson = {}

    def save_lesson():
        nonlocal current_activity, current_lesson
        if current_activity and current_lesson.get('موضوع'):
            name = normalize_name(current_activity)
            if name not in lessons:
                lessons[name] = []
            lessons[name].append(current_lesson.copy())

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue

        # نشاط جديد؟
        m = RE_ACTIVITY.search(text)
        if m:
            save_lesson()
            current_activity = m.group(1).strip()
            current_lesson = {}
            continue

        # موضوع؟
        m = RE_TOPIC.search(text)
        if m:
            current_lesson['موضوع'] = clean_text(m.group(1))
            continue

        # مؤشر كفاءة؟
        m = RE_INDICATOR.search(text)
        if m:
            current_lesson['كفاءة'] = clean_text(m.group(1))
            continue

    save_lesson()
    return lessons


def extract_from_tables(doc, existing=None):
    """استخراج الدروس من جداول المستند"""
    if existing is None:
        existing = {}

    for table in doc.tables:
        current_activity = None
        current_lesson = {}

        for row in table.rows:
            row_text = " | ".join(
                clean_text(cell.text) for cell in row.cells
            )
            for cell in row.cells:
                text = clean_text(cell.text)
                if not text:
                    continue

                m = RE_ACTIVITY.search(text)
                if m:
                    if current_activity and current_lesson.get('موضوع'):
                        name = normalize_name(current_activity)
                        existing.setdefault(name, []).append(
                            current_lesson.copy()
                        )
                    current_activity = m.group(1).strip()
                    current_lesson = {}
                    continue

                m = RE_TOPIC.search(text)
                if m:
                    current_lesson['موضوع'] = clean_text(m.group(1))
                    continue

                m = RE_INDICATOR.search(text)
                if m:
                    current_lesson['كفاءة'] = clean_text(m.group(1))
                    continue

        # حفظ الأخير
        if current_activity and current_lesson.get('موضوع'):
            name = normalize_name(current_activity)
            existing.setdefault(name, []).append(
                current_lesson.copy()
            )

    return existing


def extract_all_lessons(file_path):
    """الدالة الشاملة: تستخرج من الفقرات والجداول معاً"""
    if not os.path.exists(file_path):
        print(f"❌ الملف غير موجود: {file_path}")
        return {}
    print(f"📖 جارٍ قراءة: {file_path}")
    doc = Document(file_path)

    # استخراج من الفقرات
    lessons = extract_from_paragraphs(doc)
    # تكميل من الجداول
    lessons = extract_from_tables(doc, lessons)
    return lessons


# ═══════════════════════════════════════════════════════
# القسم 5: محرك الحقن في القالب
# ═══════════════════════════════════════════════════════


def safe_replace(paragraph, old, new):
    """استبدال ذكي يتعامل مع Runs المتعددة"""
    full = paragraph.text
    if old not in full:
        return False

    # المحاولة 1: استبدال مباشر في كل Run
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # المحاولة 2: النص موزع على عدة Runs
    new_text = full.replace(old, new)
    if paragraph.runs:
        # نحتفظ بتنسيق أول Run
        first_run = paragraph.runs[0]
        for run in paragraph.runs:
            run.text = ""
        first_run.text = new_text
        return True

    return False


def inject_into_template(template_path, output_path, replacements):
    """حقن البيانات في القالب وحفظ الملف"""
    doc = Document(template_path)
    count = 0

    # البحث في الجداول
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in para.text:
                            safe_replace(para, key, str(val))
                            count += 1

    # البحث في الفقرات العادية
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                safe_replace(para, key, str(val))
                count += 1

    doc.save(output_path)
    return count


# ═══════════════════════════════════════════════════════
# القسم 6: بناء الكراس اليومي
# ═══════════════════════════════════════════════════════


def build_daily_planner(day, template_path, output_path, schedule, lessons_db):
    """بناء كراس يوم واحد"""
    plan = schedule.get(day)
    if not plan:
        print(f"❌ اليوم '{day}' غير موجود في التوقيت")
        return False

    if not os.path.exists(template_path):
        print(f"❌ القالب غير موجود: {template_path}")
        return False

    replacements = {"{{اليوم}}": day}
    warnings = []

    print(f"\n{'═' * 55}")
    print(f" 📅 يوم {day} — {len(plan)} حصة")
    print(f"{'═' * 55}")

    for i, session in enumerate(plan, start=1):
        activity = session["النشاط"]
        duration = session["المدة"]

        # المفاتيح
        keys = {
            'نشاط': f"{{{{نشاط_{i}}}}}",
            'موضوع': f"{{{{موضوع_{i}}}}}",
            'كفاءة': f"{{{{كفاءة_{i}}}}}",
            'مدة': f"{{{{مدة_{i}}}}}",
            'ميدان': f"{{{{ميدان_{i}}}}}",
        }

        replacements[keys['مدة']] = duration
        replacements[keys['نشاط']] = activity

        if activity in ROUTINE_ACTIVITIES:
            replacements[keys['موضوع']] = "—"
            replacements[keys['كفاءة']] = "—"
            replacements[keys['ميدان']] = "—"
            print(f" ⏰ {i:2d}. {activity:<20} ({duration}) ← روتيني")
        elif activity in lessons_db and lessons_db[activity]:
            lesson = lessons_db[activity].pop(0)
            topic = lesson.get('موضوع', '—')
            indic = lesson.get('كفاءة', '—')
            replacements[keys['موضوع']] = topic
            replacements[keys['كفاءة']] = indic
            replacements[keys['ميدان']] = activity
            print(f" 📖 {i:2d}. {activity:<20} ({duration})")
            print(f" ├─ الموضوع: {topic}")
            print(f" └─ الكفاءة: {indic}")
        else:
            replacements[keys['موضوع']] = "⚠ لا توجد مذكرة"
            replacements[keys['كفاءة']] = "⚠ لا توجد مذكرة"
            replacements[keys['ميدان']] = activity
            warnings.append(activity)
            print(f" ⚠️ {i:2d}. {activity:<20} ({duration}) ← بدون مذكرة!")

    # الحقن
    count = inject_into_template(template_path, output_path, replacements)
    print(f"\n ✅ تم الحفظ: {output_path}")
    print(f" استبدالات: {count}")
    if warnings:
        print(f" ⚠ مواد بدون مذكرة: {', '.join(set(warnings))}")
    return True


# ═══════════════════════════════════════════════════════
# القسم 7: نقطة التشغيل
# ═══════════════════════════════════════════════════════


def show_summary(db):
    """عرض ملخص الدروس المستخرجة"""
    print(f"\n{'═' * 55}")
    print(" 📚 الدروس المستخرجة من المذكرات")
    print(f"{'═' * 55}")
    total = 0
    for subject in sorted(db.keys()):
        lessons = db[subject]
        total += len(lessons)
        print(f"\n 📘 {subject} ({len(lessons)} درس):")
        for j, lesson in enumerate(lessons, 1):
            t = lesson.get('موضوع', '?')
            print(f"    {j}. {t}")
    print(f"\n {'─' * 40}")
    print(f" المجموع: {total} درس في {len(db)} مادة")
    print(f"{'═' * 55}")


def show_schedule_match(db):
    """عرض مقارنة بين المواد في التوقيت والمواد المستخرجة من المذكرات"""
    # جمع كل المواد الفريدة من التوقيت
    schedule_subjects = set()
    for day_plan in weekly_schedule.values():
        for s in day_plan:
            if s["النشاط"] not in ROUTINE_ACTIVITIES:
                schedule_subjects.add(s["النشاط"])

    extracted_subjects = set(db.keys())
    matched = schedule_subjects & extracted_subjects
    missing = schedule_subjects - extracted_subjects
    extra = extracted_subjects - schedule_subjects

    print(f"\n{'─' * 55}")
    print(" 🔍 تحليل التطابق بين التوقيت والمذكرات:")
    print(f"{'─' * 55}")
    if matched:
        print(f" ✅ متطابقة ({len(matched)}):")
        for m in sorted(matched):
            print(f"    • {m}")
    if missing:
        print(f"\n ❌ موجودة في التوقيت لكن ليست في المذكرات ({len(missing)}):")
        for m in sorted(missing):
            print(f"    • {m}")
        print(" → أضف هذه المواد للمذكرات أو عدّل NAME_MAPPING")
    if extra:
        print(f"\n ℹ️ موجودة في المذكرات لكن ليست في التوقيت ({len(extra)}):")
        for m in sorted(extra):
            print(f"    • {m}")
    print(f"{'─' * 55}")


def main():
    print("""
╔══════════════════════════════════════════════════╗
║ 🎓 أتمتة الكراس اليومي - التحضيري              ║
║ الإصدار 2.0                                    ║
╚══════════════════════════════════════════════════╝
""")

    # ── الإعدادات ──
    DEFAULT_MEMO = "مذكرات_الأسبوع.docx"
    TEMPLATE = "template.docx"
    OUTPUT_DIR = "output"
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # ── 1. اسم ملف المذكرات ──
    memo = input(f"📁 ملف المذكرات [{DEFAULT_MEMO}]: ").strip()
    if not memo:
        memo = DEFAULT_MEMO

    # ── 2. استخراج الدروس ──
    lessons_db = extract_all_lessons(memo)
    if not lessons_db:
        print("\n❌ لم يتم استخراج أي درس!")
        print("تأكد من وجود الأنماط التالية في المذكرات:")
        print("    النشاط : ...")
        print("    الموضوع : ...")
        print("    مؤشر الكفاءة : ...")
        return

    show_summary(lessons_db)
    show_schedule_match(lessons_db)

    # ── 3. اختيار اليوم ──
    days = list(weekly_schedule.keys())
    print(f"\nالأيام: {' | '.join(days)} | الكل")
    choice = input("\n📅 اختر اليوم: ").strip()

    if choice == "الكل":
        selected = days
    elif choice in days:
        selected = [choice]
    else:
        print(f"❌ '{choice}' غير صحيح!")
        return

    # ── 4. نسخة عمل (لعدم تعديل الأصل) ──
    working_db = copy.deepcopy(lessons_db)

    # ── 5. التوليد ──
    ok = 0
    for day in selected:
        out = os.path.join(OUTPUT_DIR, f"كراس_{day}.docx")
        if build_daily_planner(
            day,
            TEMPLATE,
            out,
            weekly_schedule,
            working_db
        ):
            ok += 1

    # ── 6. التقرير النهائي ──
    print(f"\n{'═' * 55}")
    print(f" 🏁 تم إنشاء {ok}/{len(selected)} كراس")
    print(f" 📂 المجلد: {os.path.abspath(OUTPUT_DIR)}")
    print(f"{'═' * 55}")

    # دروس متبقية لم تُستعمل
    leftover = {k: v for k, v in working_db.items() if v}
    if leftover:
        print("\n📌 دروس لم تُستخدم:")
        for subj, lst in leftover.items():
            print(f" • {subj}: {len(lst)} درس")


if __name__ == "__main__":
    main()
