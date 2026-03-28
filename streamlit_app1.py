# -*- coding: utf-8 -*-
"""
🎓 الكراس اليومي — قسم التحضيري
الإصدار 5.0 — مدمج مع Groq AI
"""

import streamlit as st
import re, copy, json, math
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False


# ╔══════════════════════════════════════════════════════════════╗
# ║                     القسم 1: الثوابت                        ║
# ╚══════════════════════════════════════════════════════════════╝

ROUTINE_ACTIVITIES = ["الاستقبال", "الاستراحة", "تهيئة الخروج", "نهاية الخروج"]

DOMAIN_MAPPING = {
    "تعبير شفوي": "المجال اللغوي",
    "مبادئ القراءة": "المجال اللغوي",
    "تخطيط": "المجال اللغوي",
    "رياضيات": "المجال الرياضي",
    "ت علمية وتكنولوجية": "المجال العلمي",
    "ت إسلامية": "المجال الاجتماعي",
    "ت مدنية": "المجال الاجتماعي",
    "تربية تشكيلية": "المجال الفني",
    "موسيقى وإنشاد": "المجال الفني",
    "مسرح وعرائس": "المجال الفني",
    "ت بدنية": "المجال البدني والإيقاعي",
    "ت إيقاعية": "المجال البدني والإيقاعي",
}

SUBJECT_WEEKLY_COUNT = {
    "تعبير شفوي": 3, "مبادئ القراءة": 4, "تخطيط": 2,
    "رياضيات": 5, "ت علمية وتكنولوجية": 4,
    "ت إسلامية": 2, "ت مدنية": 2,
    "تربية تشكيلية": 2, "موسيقى وإنشاد": 2, "مسرح وعرائس": 2,
    "ت بدنية": 4, "ت إيقاعية": 2,
}

# الأسماء الموحدة التي يجب أن يُرجعها الذكاء الاصطناعي
CANONICAL_NAMES = list(SUBJECT_WEEKLY_COUNT.keys())

DOMAIN_COLORS = {
    "المجال اللغوي": "#1565C0",
    "المجال الرياضي": "#C62828",
    "المجال العلمي": "#2E7D32",
    "المجال الاجتماعي": "#F57F17",
    "المجال الفني": "#6A1B9A",
    "المجال البدني والإيقاعي": "#00838F",
}

NAME_MAPPING = {
    "تربية علمية": "ت علمية وتكنولوجية",
    "تربية تكنولوجية": "ت علمية وتكنولوجية",
    "تربية علمية وتكنولوجية": "ت علمية وتكنولوجية",
    "ت علمية": "ت علمية وتكنولوجية",
    "ت تكنولوجية": "ت علمية وتكنولوجية",
    "علوم وتكنولوجيا": "ت علمية وتكنولوجية",
    "علمية وتكنولوجية": "ت علمية وتكنولوجية",
    "علمية": "ت علمية وتكنولوجية",
    "تكنولوجية": "ت علمية وتكنولوجية",
    "التربية العلمية والتكنولوجية": "ت علمية وتكنولوجية",
    "التربية العلمية": "ت علمية وتكنولوجية",
    "التربية التكنولوجية": "ت علمية وتكنولوجية",
    "تعبير": "تعبير شفوي", "التعبير الشفوي": "تعبير شفوي",
    "التعبير": "تعبير شفوي", "تعبير شفهي": "تعبير شفوي",
    "شفوي": "تعبير شفوي",
    "قراءة": "مبادئ القراءة", "القراءة": "مبادئ القراءة",
    "مبادئ في القراءة": "مبادئ القراءة",
    "مبادىء القراءة": "مبادئ القراءة",
    "كتابة": "تخطيط", "خط": "تخطيط", "الكتابة": "تخطيط",
    "مبادئ التخطيط": "تخطيط", "مبادئ الكتابة": "تخطيط",
    "الرياضيات": "رياضيات",
    "تربية إسلامية": "ت إسلامية", "التربية الإسلامية": "ت إسلامية",
    "إسلامية": "ت إسلامية", "تربية اسلامية": "ت إسلامية",
    "تربية مدنية": "ت مدنية", "التربية المدنية": "ت مدنية",
    "مدنية": "ت مدنية",
    "رسم": "تربية تشكيلية", "رسم وأشغال": "تربية تشكيلية",
    "أشغال يدوية": "تربية تشكيلية", "تربية فنية": "تربية تشكيلية",
    "فنون تشكيلية": "تربية تشكيلية", "التربية التشكيلية": "تربية تشكيلية",
    "تشكيلية": "تربية تشكيلية",
    "موسيقى": "موسيقى وإنشاد", "إنشاد": "موسيقى وإنشاد",
    "تربية موسيقية": "موسيقى وإنشاد", "الموسيقى": "موسيقى وإنشاد",
    "الموسيقى والإنشاد": "موسيقى وإنشاد",
    "مسرح": "مسرح وعرائس", "المسرح": "مسرح وعرائس",
    "المسرح والعرائس": "مسرح وعرائس",
    "تربية بدنية": "ت بدنية", "التربية البدنية": "ت بدنية",
    "بدنية": "ت بدنية", "تربية بدنية ورياضية": "ت بدنية",
    "تربية إيقاعية": "ت إيقاعية", "التربية الإيقاعية": "ت إيقاعية",
    "إيقاعية": "ت إيقاعية", "تربية ايقاعية": "ت إيقاعية",
}

weekly_schedule = {
    "الأحد": [
        {"النشاط": "الاستقبال", "المدة": "15 د", "الفترة": "صباحية"},
        {"النشاط": "تعبير شفوي", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "مبادئ القراءة", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "رياضيات", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت علمية وتكنولوجية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت إسلامية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د", "الفترة": "صباحية"},
        {"النشاط": "الاستقبال", "المدة": "15 د", "الفترة": "مسائية"},
        {"النشاط": "مسرح وعرائس", "المدة": "30 د", "الفترة": "مسائية"},
        {"النشاط": "تربية تشكيلية", "المدة": "30 د", "الفترة": "مسائية"},
        {"النشاط": "ت بدنية", "المدة": "30 د", "الفترة": "مسائية"},
        {"النشاط": "نهاية الخروج", "المدة": "15 د", "الفترة": "مسائية"},
    ],
    "الإثنين": [
        {"النشاط": "الاستقبال", "المدة": "15 د", "الفترة": "صباحية"},
        {"النشاط": "رياضيات", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "تعبير شفوي", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "تخطيط", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت علمية وتكنولوجية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت مدنية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د", "الفترة": "صباحية"},
        {"النشاط": "الاستقبال", "المدة": "15 د", "الفترة": "مسائية"},
        {"النشاط": "مسرح وعرائس", "المدة": "30 د", "الفترة": "مسائية"},
        {"النشاط": "تربية تشكيلية", "المدة": "30 د", "الفترة": "مسائية"},
        {"النشاط": "ت بدنية", "المدة": "30 د", "الفترة": "مسائية"},
        {"النشاط": "نهاية الخروج", "المدة": "15 د", "الفترة": "مسائية"},
    ],
    "الثلاثاء": [
        {"النشاط": "الاستقبال", "المدة": "15 د", "الفترة": "صباحية"},
        {"النشاط": "تعبير شفوي", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "مبادئ القراءة", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "رياضيات", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت إسلامية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت بدنية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د", "الفترة": "صباحية"},
    ],
    "الأربعاء": [
        {"النشاط": "الاستقبال", "المدة": "15 د", "الفترة": "صباحية"},
        {"النشاط": "رياضيات", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "مبادئ القراءة", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "تخطيط", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت علمية وتكنولوجية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت مدنية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د", "الفترة": "صباحية"},
        {"النشاط": "الاستقبال", "المدة": "15 د", "الفترة": "مسائية"},
        {"النشاط": "ت إيقاعية", "المدة": "30 د", "الفترة": "مسائية"},
        {"النشاط": "موسيقى وإنشاد", "المدة": "30 د", "الفترة": "مسائية"},
        {"النشاط": "ت بدنية", "المدة": "30 د", "الفترة": "مسائية"},
        {"النشاط": "نهاية الخروج", "المدة": "15 د", "الفترة": "مسائية"},
    ],
    "الخميس": [
        {"النشاط": "الاستقبال", "المدة": "15 د", "الفترة": "صباحية"},
        {"النشاط": "مبادئ القراءة", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "رياضيات", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت علمية وتكنولوجية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "ت إيقاعية", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "موسيقى وإنشاد", "المدة": "30 د", "الفترة": "صباحية"},
        {"النشاط": "تهيئة الخروج", "المدة": "15 د", "الفترة": "صباحية"},
    ],
}


# ╔══════════════════════════════════════════════════════════════╗
# ║                  القسم 2: دوال مساعدة                       ║
# ╚══════════════════════════════════════════════════════════════╝

def get_all_teaching_subjects():
    s = set()
    for p in weekly_schedule.values():
        for x in p:
            if x["النشاط"] not in ROUTINE_ACTIVITIES:
                s.add(x["النشاط"])
    return s

def get_domain_for(act):
    return DOMAIN_MAPPING.get(act, "—")

def verify_schedule():
    c = {}
    for p in weekly_schedule.values():
        for x in p:
            a = x["النشاط"]
            if a not in ROUTINE_ACTIVITIES:
                c[a] = c.get(a, 0) + 1
    errs = []
    for s, exp in SUBJECT_WEEKLY_COUNT.items():
        if c.get(s, 0) != exp:
            errs.append(f"{s}: متوقع {exp} | فعلي {c.get(s, 0)}")
    return errs, c

def domain_badge(d):
    col = DOMAIN_COLORS.get(d, "#666")
    return (f'<span style="display:inline-block;padding:2px 10px;border-radius:12px;'
            f'font-size:0.75rem;font-weight:600;background:{col}22;color:{col};'
            f'border:1px solid {col}44;">{d}</span>')

def clean_text(text):
    if not text: return ""
    text = re.sub(r'ـ+', '', text)
    text = re.sub(r'[\u0610-\u061A\u064B-\u065F\u0670\u06D6-\u06DC\u06DF-\u06E4\u06E7\u06E8\u06EA-\u06ED]', '', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[\.\s]+$', '', text)
    return text.strip()

def normalize_name(raw):
    c = clean_text(raw)
    if not c: return c
    if c in NAME_MAPPING: return NAME_MAPPING[c]
    for k, v in NAME_MAPPING.items():
        if k in c or c in k: return v
    n = c.replace('أ','ا').replace('إ','ا').replace('آ','ا').replace('ى','ي').replace('ة','ه')
    for k, v in NAME_MAPPING.items():
        kn = k.replace('أ','ا').replace('إ','ا').replace('آ','ا').replace('ى','ي').replace('ة','ه')
        if kn == n or kn in n or n in kn: return v
    roots = {"علم":"ت علمية وتكنولوجية","تكنولوج":"ت علمية وتكنولوجية",
             "إسلام":"ت إسلامية","اسلام":"ت إسلامية","مدن":"ت مدنية",
             "بدن":"ت بدنية","إيقاع":"ت إيقاعية","ايقاع":"ت إيقاعية",
             "تشكيل":"تربية تشكيلية","رياض":"رياضيات","قراء":"مبادئ القراءة",
             "تخطيط":"تخطيط","كتاب":"تخطيط","شفو":"تعبير شفوي",
             "تعبير":"تعبير شفوي","مسرح":"مسرح وعرائس","عرائس":"مسرح وعرائس",
             "موسيق":"موسيقى وإنشاد","إنشاد":"موسيقى وإنشاد","انشاد":"موسيقى وإنشاد",
             "رسم":"تربية تشكيلية","أشغال":"تربية تشكيلية"}
    for r, t in roots.items():
        if r in c: return t
    return c


# ╔══════════════════════════════════════════════════════════════╗
# ║          القسم 3: قراءة محتوى ملف المذكرات                  ║
# ╚══════════════════════════════════════════════════════════════╝

def read_docx_full_text(file_bytes):
    """قراءة كل النص من ملف Word (فقرات + جداول)"""
    doc = Document(BytesIO(file_bytes))
    lines = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    row_texts.append(t)
            if row_texts:
                lines.append(" | ".join(row_texts))

    return "\n".join(lines)


# ╔══════════════════════════════════════════════════════════════╗
# ║          القسم 4: الاستخراج بالذكاء الاصطناعي (Groq)       ║
# ╚══════════════════════════════════════════════════════════════╝

AI_PROMPT = """أنت محلل وثائق تربوية متخصص في المنهاج الجزائري لقسم التحضيري.

سأعطيك محتوى ملف مذكرات أسبوعية. مهمتك استخراج كل الدروس.

المواد المتوقعة بأسمائها الرسمية (استخدم هذه الأسماء بالضبط):
- تعبير شفوي
- مبادئ القراءة
- تخطيط
- رياضيات
- ت علمية وتكنولوجية  (تشمل التربية العلمية والتكنولوجية معاً)
- ت إسلامية
- ت مدنية
- تربية تشكيلية  (تشمل الرسم والأشغال)
- موسيقى وإنشاد
- مسرح وعرائس
- ت بدنية
- ت إيقاعية

لكل درس استخرج:
1. "مادة": اسم المادة (استخدم الأسماء الرسمية أعلاه بالضبط)
2. "موضوع": عنوان/موضوع الدرس
3. "كفاءة": مؤشر الكفاءة المستهدفة

تعليمات مهمة:
- إذا كانت التربية العلمية والتربية التكنولوجية مذكورتين بشكل منفصل، اجمعهما تحت "ت علمية وتكنولوجية"
- إذا وجدت أكثر من درس لنفس المادة، أضف كل واحد كعنصر مستقل
- إذا لم تجد مؤشر كفاءة واضح، اكتب "—"
- أعد النتيجة كـ JSON فقط بدون أي نص إضافي

الشكل المطلوب:
[
  {"مادة": "تعبير شفوي", "موضوع": "عنوان الدرس هنا", "كفاءة": "مؤشر الكفاءة هنا"},
  {"مادة": "رياضيات", "موضوع": "...", "كفاءة": "..."}
]

محتوى المذكرات:
---
{content}
---"""


def parse_ai_json(response_text):
    """تحليل رد الذكاء الاصطناعي واستخراج JSON"""
    # محاولة 1: تحليل مباشر
    try:
        return json.loads(response_text)
    except (json.JSONDecodeError, TypeError):
        pass

    # محاولة 2: استخراج من code block
    m = re.search(r'```(?:json)?\s*([\s\S]*?)```', response_text)
    if m:
        try:
            return json.loads(m.group(1))
        except (json.JSONDecodeError, TypeError):
            pass

    # محاولة 3: إيجاد مصفوفة JSON
    m = re.search(r'\[[\s\S]*\]', response_text)
    if m:
        try:
            return json.loads(m.group(0))
        except (json.JSONDecodeError, TypeError):
            pass

    return None


def extract_with_ai(text_content, api_key, model="llama-3.3-70b-versatile"):
    """
    استخراج الدروس باستخدام Groq AI
    المُخرج: (lessons_db, raw_response, error)
    """
    if not GROQ_AVAILABLE:
        return None, None, "مكتبة groq غير مثبتة. شغّل: pip install groq"

    if not api_key:
        return None, None, "مفتاح Groq API مطلوب"

    # تقليص النص إذا كان طويلاً جداً
    max_chars = 12000
    if len(text_content) > max_chars:
        text_content = text_content[:max_chars] + "\n... (تم اختصار النص)"

    prompt = AI_PROMPT.replace("{content}", text_content)

    try:
        client = Groq(api_key=api_key)
        completion = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": "أنت محلل وثائق تربوية. أجب بـ JSON فقط."
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=4000,
        )

        raw = completion.choices[0].message.content
        parsed = parse_ai_json(raw)

        if not parsed or not isinstance(parsed, list):
            return None, raw, "فشل تحليل JSON من رد الذكاء الاصطناعي"

        # تحويل إلى lessons_db
        lessons_db = {}
        for item in parsed:
            if not isinstance(item, dict):
                continue
            subject = item.get("مادة", "").strip()
            topic = item.get("موضوع", "—").strip()
            indicator = item.get("كفاءة", "—").strip()

            if not subject or not topic:
                continue

            # توحيد الاسم
            normalized = normalize_name(subject)
            # التحقق إذا كان اسماً معروفاً
            if normalized not in SUBJECT_WEEKLY_COUNT:
                # محاولة أخيرة
                normalized = subject

            lessons_db.setdefault(normalized, []).append({
                "موضوع": topic,
                "كفاءة": indicator,
            })

        return lessons_db, raw, None

    except Exception as e:
        return None, None, f"خطأ في Groq API: {str(e)}"


# ╔══════════════════════════════════════════════════════════════╗
# ║          القسم 5: الاستخراج بـ Regex (احتياطي)              ║
# ╚══════════════════════════════════════════════════════════════╝

RE_ACT = re.compile(r'(?:النشاط|المادة|نشاط|مادة|المجال|مجال\s*التعلم?|الميدان)\s*[:/\-\.\|]*\s*(.+)')
RE_TOP = re.compile(r'(?:الموضوع|موضوع|الوحدة|وحدة|عنوان\s*الدرس|المحتوى)\s*[:/\-\.\|]*\s*(.+)')
RE_IND = re.compile(r'(?:مؤشر\s*(?:ات)?\s*الكفاء\w*|الكفاءة\s*المستهدفة|كفاءة)\s*[:/\-\.\|]*\s*(.+)')

def extract_with_regex(file_bytes):
    doc = Document(BytesIO(file_bytes))
    lessons = {}
    cur_act = None
    cur_les = {}

    def _save():
        nonlocal cur_act, cur_les
        if cur_act and cur_les.get('موضوع'):
            name = normalize_name(cur_act)
            lessons.setdefault(name, []).append(cur_les.copy())

    def _try(text):
        nonlocal cur_act, cur_les
        m = RE_ACT.search(text)
        if m:
            _save()
            cur_act = m.group(1).strip()
            cur_les = {}
            return
        m = RE_TOP.search(text)
        if m:
            cur_les['موضوع'] = clean_text(m.group(1))
            return
        m = RE_IND.search(text)
        if m:
            cur_les['كفاءة'] = clean_text(m.group(1))

    for para in doc.paragraphs:
        t = clean_text(para.text)
        if t: _try(t)
    _save()

    cur_act = None
    cur_les = {}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for line in cell.text.split('\n'):
                    t = clean_text(line)
                    if t: _try(t)
    _save()
    return lessons


# ╔══════════════════════════════════════════════════════════════╗
# ║         القسم 6: محرك التوزيع الذكي للحصص                  ║
# ╚══════════════════════════════════════════════════════════════╝

def distribute_lessons(raw_db):
    """
    توزيع الدروس على حصص الأسبوع بذكاء

    المنطق:
    - إذا مادة لها 5 حصص ودرس واحد → نفس الدرس يتكرر 5 مرات
    - إذا مادة لها 5 حصص ودرسان → الأول 3 حصص والثاني 2
    - إذا مادة لها 5 حصص و5 دروس → درس لكل حصة

    المُخرج: قاموس بنفس الشكل لكن القوائم مملوءة حسب عدد الحصص
    """
    distributed = {}
    distribution_report = {}

    for subject, required in SUBJECT_WEEKLY_COUNT.items():
        available = raw_db.get(subject, [])
        count = len(available)

        if count == 0:
            distributed[subject] = []
            distribution_report[subject] = {
                "مطلوب": required, "متوفر": 0,
                "حالة": "❌ لا يوجد", "توزيع": []
            }
            continue

        result = []

        if count >= required:
            # دروس كافية أو أكثر: نأخذ أول N
            result = [les.copy() for les in available[:required]]
            dist_detail = [f"درس {i+1}" for i in range(required)]

        else:
            # دروس أقل من الحصص: نوزع بالتساوي
            per_lesson = required / count
            dist_detail = []

            for i, lesson in enumerate(available):
                start = round(i * per_lesson)
                end = round((i + 1) * per_lesson)
                num_slots = end - start

                for slot in range(num_slots):
                    enriched = lesson.copy()
                    enriched["_رقم_الحصة"] = len(result) + 1
                    enriched["_إجمالي_الحصص"] = required
                    result.append(enriched)
                    dist_detail.append(f"درس {i+1}")

        distributed[subject] = result
        distribution_report[subject] = {
            "مطلوب": required,
            "متوفر": count,
            "حالة": "✅" if result else "❌",
            "توزيع": dist_detail
        }

    return distributed, distribution_report


# ╔══════════════════════════════════════════════════════════════╗
# ║              القسم 7: بناء القالب                            ║
# ╚══════════════════════════════════════════════════════════════╝

def _rtl(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn('w:bidi'), {}))

def _cell(c, t, bold=False, size=10, color=None):
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl(p)
    r = p.add_run(t)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Sakkal Majalla"
    if color: r.font.color.rgb = color
    rPr = r._r.get_or_add_rPr()
    rPr.append(rPr.makeelement(qn('w:rFonts'), {qn('w:cs'): 'Sakkal Majalla'}))

def _shade(row, hx):
    for c in row.cells:
        tc = c._tc.get_or_add_tcPr()
        tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): hx, qn('w:val'): 'clear'}))

def _ptable(doc, title, start, count):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl(h)
    r = h.add_run(title)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 51, 102)
    hdrs = ['مؤشرات الكفاءة','عنوان الدرس','الميدان','النشاط','المدة']
    tbl = doc.add_table(rows=1+count, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ws = [Cm(5.5),Cm(5),Cm(3.5),Cm(3.5),Cm(2)]
    for row in tbl.rows:
        for i,w in enumerate(ws): row.cells[i].width = w
    hdr = tbl.rows[0]
    _shade(hdr, "1F4E79")
    for i,t in enumerate(hdrs):
        _cell(hdr.cells[i], t, True, 11, RGBColor(255,255,255))
    for j in range(count):
        n = start+j
        dr = tbl.rows[1+j]
        if j%2==0: _shade(dr,"EDF2F9")
        for i,ph in enumerate([f'{{{{كفاءة_{n}}}}}',f'{{{{موضوع_{n}}}}}',f'{{{{ميدان_{n}}}}}',f'{{{{نشاط_{n}}}}}',f'{{{{مدة_{n}}}}}']):
            _cell(dr.cells[i], ph, size=9)

def create_template_bytes():
    doc = Document()
    for sec in doc.sections:
        sec._sectPr.append(sec._sectPr.makeelement(qn('w:bidi'), {}))
    for t,sz,b in [('الجمهورية الجزائرية الديمقراطية الشعبية',12,True),('وزارة التربية الوطنية',11,False)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _rtl(p)
        r = p.add_run(t); r.bold = b; r.font.size = Pt(sz)
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER; _rtl(tp)
    tr = tp.add_run('الكراس اليومي'); tr.bold = True; tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0,51,102)
    info = doc.add_table(rows=1, cols=3)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell(info.rows[0].cells[2], 'اليوم : {{اليوم}}', True, 12)
    _cell(info.rows[0].cells[1], 'التاريخ : {{التاريخ}}', size=11)
    _cell(info.rows[0].cells[0], 'الأسبوع : {{الأسبوع}}', size=11)
    doc.add_paragraph('')
    _ptable(doc, '☀ الفترة الصباحية', 1, 7)
    doc.add_paragraph('')
    _ptable(doc, '🌙 الفترة المسائية', 8, 5)
    doc.add_paragraph('')
    np2 = doc.add_paragraph()
    np2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; _rtl(np2)
    np2.add_run('ملاحظات : ' + '.'*80).font.size = Pt(10)
    buf = BytesIO(); doc.save(buf)
    return buf.getvalue()


# ╔══════════════════════════════════════════════════════════════╗
# ║               القسم 8: محرك الحقن والتوليد                  ║
# ╚══════════════════════════════════════════════════════════════╝

def _safe_replace(para, old, new):
    if old not in para.text: return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    full = para.text.replace(old, new)
    if para.runs:
        for run in para.runs: run.text = ""
        para.runs[0].text = full
    return True

def build_daily_planner(day, template_bytes, distributed_db, week_num="", date_str=""):
    plan = weekly_schedule.get(day, [])
    if not plan:
        return None, [], [f"اليوم '{day}' غير موجود"]
    doc = Document(BytesIO(template_bytes))
    reps = {"{{اليوم}}": day, "{{التاريخ}}": date_str, "{{الأسبوع}}": week_num}
    sessions_info = []
    warnings = []

    for i, session in enumerate(plan, 1):
        act = session["النشاط"]
        dur = session["المدة"]
        per = session.get("الفترة", "")
        domain = DOMAIN_MAPPING.get(act, "—")
        ka = f"{{{{نشاط_{i}}}}}"
        kt = f"{{{{موضوع_{i}}}}}"
        ki = f"{{{{كفاءة_{i}}}}}"
        kd = f"{{{{مدة_{i}}}}}"
        kf = f"{{{{ميدان_{i}}}}}"
        reps[kd] = dur; reps[ka] = act
        info = {"رقم":i,"النشاط":act,"المدة":dur,"الفترة":per,"المجال":domain,
                "نوع":"روتيني","الموضوع":"—","الكفاءة":"—"}

        if act in ROUTINE_ACTIVITIES:
            reps[kt]="—"; reps[ki]="—"; reps[kf]="—"
        elif act in distributed_db and distributed_db[act]:
            lesson = distributed_db[act].pop(0)
            topic = lesson.get('موضوع','—')
            indic = lesson.get('كفاءة','—')
            sess_num = lesson.get('_رقم_الحصة','')
            total_s = lesson.get('_إجمالي_الحصص','')
            reps[kt]=topic; reps[ki]=indic; reps[kf]=domain
            info["نوع"]="تعليمي"; info["الموضوع"]=topic; info["الكفاءة"]=indic
            if sess_num:
                info["رقم_الحصة"] = f"{sess_num}/{total_s}"
        else:
            reps[kt]="⚠ لا توجد مذكرة"; reps[ki]="⚠ لا توجد مذكرة"
            reps[kf]=domain; info["نوع"]="ناقص"; warnings.append(act)
        sessions_info.append(info)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for k,v in reps.items():
                        if k in para.text: _safe_replace(para,k,str(v))
    for para in doc.paragraphs:
        for k,v in reps.items():
            if k in para.text: _safe_replace(para,k,str(v))
    buf = BytesIO(); doc.save(buf)
    return buf.getvalue(), sessions_info, warnings


# ╔══════════════════════════════════════════════════════════════╗
# ║                  القسم 9: واجهة Streamlit                   ║
# ╚══════════════════════════════════════════════════════════════╝

st.set_page_config(page_title="الكراس اليومي 🎓", page_icon="🎓",
                   layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
.main .block-container{direction:rtl;text-align:right}
h1,h2,h3{text-align:center!important}
.card{padding:1rem;border-radius:12px;text-align:center;margin:.4rem 0;box-shadow:0 2px 8px rgba(0,0,0,.08)}
.card h4{margin:0 0 .3rem 0;font-size:.9rem}
.card .num{font-size:2rem;font-weight:700}
.card-blue{background:linear-gradient(135deg,#1F4E79,#2E75B6);color:#fff}
.card-green{background:linear-gradient(135deg,#2E7D32,#43A047);color:#fff}
.card-amber{background:linear-gradient(135deg,#E65100,#FF9800);color:#fff}
.card-purple{background:linear-gradient(135deg,#4A148C,#7B1FA2);color:#fff}
.slot{display:flex;align-items:center;gap:.8rem;padding:.7rem 1rem;margin:.3rem 0;border-radius:8px;direction:rtl}
.slot-routine{background:#f5f5f5;border-right:4px solid #9e9e9e}
.slot-teach{background:#e3f2fd;border-right:4px solid #1565c0}
.slot-warn{background:#fff3e0;border-right:4px solid #e65100}
.stDownloadButton>button{width:100%;background:linear-gradient(135deg,#1F4E79,#2E75B6)!important;color:#fff!important;border:none!important;border-radius:8px!important}
[data-testid="stSidebar"]{direction:rtl;text-align:right}
footer{visibility:hidden}
.ok-box{background:#e8f5e9;border:1px solid #4caf50;border-radius:10px;padding:1rem;text-align:center}
.ai-box{background:linear-gradient(135deg,#E8EAF6,#C5CAE9);border:2px solid #3F51B5;border-radius:12px;padding:1.2rem;margin:.5rem 0}
.dist-row{display:flex;align-items:center;gap:6px;margin:3px 0;direction:rtl}
.dist-slot{width:28px;height:28px;border-radius:6px;display:flex;align-items:center;justify-content:center;font-size:.7rem;font-weight:700;color:#fff}
</style>""", unsafe_allow_html=True)

# Session State
for k in ['lessons_db','template_bytes','generated_files','dist_report',
          'ai_raw','extraction_method','file_text']:
    if k not in st.session_state:
        st.session_state[k] = None if k != 'generated_files' else {}


# ═══════════════════════════════════════
#  الشريط الجانبي
# ═══════════════════════════════════════

with st.sidebar:
    st.markdown("## ⚙️ الإعدادات")
    st.markdown("---")
    week_num = st.text_input("📅 رقم الأسبوع", placeholder="10")
    date_str = st.text_input("📆 التاريخ", placeholder="2024/12/01")

    st.markdown("---")
    st.markdown("### 🧠 Groq AI")
    groq_key = st.text_input("🔑 مفتاح API", type="password",
                              help="احصل على مفتاح مجاني من console.groq.com")
    ai_model = st.selectbox("النموذج", [
        "llama-3.3-70b-versatile",
        "llama-3.1-70b-versatile",
        "mixtral-8x7b-32768",
        "gemma2-9b-it",
    ])

    st.markdown("---")
    st.markdown("### 📤 ملف المذكرات")
    uploaded = st.file_uploader("اختر ملف .docx", type=["docx"])

    st.markdown("---")
    if st.button("🔨 إنشاء قالب", use_container_width=True):
        st.session_state.template_bytes = create_template_bytes()
        st.success("✅")
    if st.session_state.template_bytes:
        st.download_button("📄 تحميل القالب", data=st.session_state.template_bytes,
                           file_name="template.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    st.markdown("---")
    errs, cnts = verify_schedule()
    if errs:
        st.error("❌ خلل!")
        for e in errs: st.text(e)
    else:
        st.success(f"✅ التوقيت سليم ({sum(cnts.values())} حصة)")
    st.caption("🎓 v5.0 — AI Edition")


# ═══════════════════════════════════════
#  معالجة الرفع
# ═══════════════════════════════════════

if uploaded:
    file_bytes = uploaded.read()
    if st.session_state.get('_last') != uploaded.name:
        st.session_state.file_text = read_docx_full_text(file_bytes)
        st.session_state._file_bytes = file_bytes
        st.session_state._last = uploaded.name
        st.session_state.lessons_db = None
        st.session_state.generated_files = {}
        st.session_state.dist_report = None
        st.session_state.ai_raw = None
        st.session_state.extraction_method = None
    if not st.session_state.template_bytes:
        st.session_state.template_bytes = create_template_bytes()


# ═══════════════════════════════════════
#  العنوان
# ═══════════════════════════════════════

st.markdown("""
<h1 style="background:linear-gradient(135deg,#1F4E79,#2E75B6);
-webkit-background-clip:text;-webkit-text-fill-color:transparent;
font-size:2.5rem;">🎓🧠 الكراس اليومي الذكي</h1>
<p style="text-align:center;color:#888;">مدعوم بالذكاء الاصطناعي — Groq AI</p>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════
#  التبويبات
# ═══════════════════════════════════════

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "🧠 الاستخراج", "📅 توليد", "👁️ معاينة",
    "📊 التوزيع", "🗺️ المجالات", "📄 محتوى الملف"
])


# ──── تبويب 1: الاستخراج ────

with tab1:
    if not st.session_state.file_text:
        st.info("👆 ارفع ملف المذكرات من الشريط الجانبي")
        with st.expander("📖 كيف يعمل؟"):
            st.markdown("""
            ### الطريقة الجديدة مع AI:
            1. ارفع ملف المذكرات
            2. اضغط **🧠 استخراج بالذكاء الاصطناعي**
            3. الذكاء الاصطناعي يقرأ الملف ويستخرج كل الدروس تلقائياً
            4. يتم توزيع الدروس على حصص الأسبوع بذكاء
            5. اذهب لتبويب **توليد** وحمّل الكراسات!

            ### بدون AI (احتياطي):
            يمكنك استخدام الاستخراج بالأنماط النصية (Regex)
            """)
    else:
        st.markdown("### 🎯 اختر طريقة الاستخراج")

        col_ai, col_regex = st.columns(2)

        with col_ai:
            st.markdown("""
            <div class="ai-box">
                <h3 style="text-align:center;color:#283593;">🧠 ذكاء اصطناعي</h3>
                <p style="text-align:center;">يفهم أي تنسيق ويستخرج بدقة عالية</p>
            </div>
            """, unsafe_allow_html=True)

            ai_disabled = not groq_key or not GROQ_AVAILABLE

            if ai_disabled and not GROQ_AVAILABLE:
                st.error("مكتبة groq غير مثبتة")
            elif ai_disabled:
                st.warning("أدخل مفتاح Groq API في الشريط الجانبي")

            if st.button("🧠 استخراج بالذكاء الاصطناعي",
                         use_container_width=True, type="primary",
                         disabled=ai_disabled):
                with st.spinner("🧠 الذكاء الاصطناعي يقرأ المذكرات..."):
                    db, raw, error = extract_with_ai(
                        st.session_state.file_text, groq_key, ai_model
                    )
                    st.session_state.ai_raw = raw

                if error:
                    st.error(f"❌ {error}")
                elif db:
                    # توزيع ذكي
                    distributed, report = distribute_lessons(db)
                    st.session_state.lessons_db = distributed
                    st.session_state.dist_report = report
                    st.session_state.extraction_method = "🧠 AI"
                    st.session_state.generated_files = {}
                    st.success(f"✅ تم استخراج {sum(len(v) for v in db.values())} درس وتوزيعها!")
                    st.rerun()

        with col_regex:
            st.markdown("""
            <div style="background:#F5F5F5;border:2px solid #9E9E9E;
            border-radius:12px;padding:1.2rem;margin:.5rem 0;">
                <h3 style="text-align:center;color:#616161;">📝 أنماط نصية</h3>
                <p style="text-align:center;">يبحث عن كلمات مفتاحية محددة</p>
            </div>
            """, unsafe_allow_html=True)

            if st.button("📝 استخراج بالأنماط", use_container_width=True):
                with st.spinner("📝 جارٍ البحث..."):
                    db = extract_with_regex(st.session_state._file_bytes)

                if db:
                    distributed, report = distribute_lessons(db)
                    st.session_state.lessons_db = distributed
                    st.session_state.dist_report = report
                    st.session_state.extraction_method = "📝 Regex"
                    st.session_state.generated_files = {}
                    st.success(f"✅ تم استخراج {sum(len(v) for v in db.values())} درس!")
                    st.rerun()
                else:
                    st.error("❌ لم يتم العثور على دروس!")

        # ── عرض النتائج ──
        db = st.session_state.lessons_db
        if db:
            st.markdown("---")
            method = st.session_state.extraction_method or ""
            total = sum(len(v) for v in db.values())
            subjects = get_all_teaching_subjects()
            matched = subjects & set(k for k,v in db.items() if v)
            missing = subjects - set(k for k,v in db.items() if v)

            st.markdown(f"### {method} النتائج")

            c1,c2,c3,c4 = st.columns(4)
            for col,title,num,cls in [
                (c1,"📘 المواد",len([v for v in db.values() if v]),"card-blue"),
                (c2,"📖 الحصص",total,"card-green"),
                (c3,"✅ مغطاة",len(matched),"card-purple"),
                (c4,"⚠ ناقصة",len(missing),"card-amber"),
            ]:
                with col:
                    st.markdown(f'<div class="card {cls}"><h4>{title}</h4>'
                                f'<div class="num">{num}</div></div>',
                                unsafe_allow_html=True)

            if missing:
                st.warning(f"⚠️ مواد ناقصة: **{' ، '.join(missing)}**")

            st.markdown("---")
            for subj in sorted(db.keys()):
                lessons = db[subj]
                if not lessons: continue
                domain = get_domain_for(subj)
                icon = "✅" if subj in subjects else "ℹ️"
                with st.expander(f"{icon} {subj} — {len(lessons)} حصة — {domain}"):
                    for j, les in enumerate(lessons, 1):
                        sn = les.get('_رقم_الحصة','')
                        ts = les.get('_إجمالي_الحصص','')
                        badge = f" (الحصة {sn}/{ts})" if sn else ""
                        st.markdown(
                            f"**حصة {j}{badge}:** 📝 {les.get('موضوع','—')}\n\n"
                            f"🎯 {les.get('كفاءة','—')}"
                        )
                        if j < len(lessons): st.divider()


# ──── تبويب 2: توليد ────

with tab2:
    db = st.session_state.lessons_db
    if not db or not any(db.values()):
        st.info("🧠 استخرج الدروس أولاً من التبويب الأول")
    else:
        st.markdown("### 📅 اختر الأيام")
        days = list(weekly_schedule.keys())
        cols = st.columns(len(days))
        selected = []
        for i,d in enumerate(days):
            plan = weekly_schedule[d]
            teach = sum(1 for s in plan if s["النشاط"] not in ROUTINE_ACTIVITIES)
            has_ev = any(s.get("الفترة")=="مسائية" for s in plan)
            label = f"{d} ({teach})" + ("" if has_ev else " ☀")
            with cols[i]:
                if st.checkbox(label, key=f"d_{d}"): selected.append(d)
        if st.checkbox("✅ الكل"): selected = days
        st.markdown("---")

        if selected and st.button(f"🚀 توليد {len(selected)} كراس",
                                   type="primary", use_container_width=True):
            tmpl = st.session_state.template_bytes or create_template_bytes()
            st.session_state.template_bytes = tmpl
            wdb = copy.deepcopy(db)
            gen = {}
            bar = st.progress(0)
            msg = st.empty()
            for idx,d in enumerate(selected):
                msg.text(f"⏳ {d}...")
                bar.progress(idx/len(selected))
                result,info,warns = build_daily_planner(d,tmpl,wdb,week_num,date_str)
                if result:
                    gen[d] = {'bytes':result,'sessions':info,'warnings':warns}
            bar.progress(1.0)
            msg.empty()
            st.session_state.generated_files = gen
            st.markdown(f'<div class="ok-box"><h3>✅ تم توليد {len(gen)} كراس!</h3></div>',
                        unsafe_allow_html=True)

        gf = st.session_state.generated_files
        if gf:
            st.markdown("### 📥 التحميل")
            dl_cols = st.columns(min(len(gf),3))
            for i,(d,data) in enumerate(gf.items()):
                with dl_cols[i%3]:
                    w = data.get('warnings',[])
                    st.download_button(
                        f"{'⚠️' if w else '📄'} كراس {d}",
                        data=data['bytes'], file_name=f"كراس_{d}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True, key=f"dl_{d}")
                    if w: st.caption(f"⚠ {', '.join(set(w))}")


# ──── تبويب 3: معاينة ────

with tab3:
    gf = st.session_state.generated_files
    if not gf:
        st.info("📅 ولّد الكراسات أولاً")
    else:
        dp = st.selectbox("اختر اليوم", list(gf.keys()), key="prev")
        if dp:
            data = gf[dp]
            sessions = data['sessions']
            st.markdown(f"## 📋 كراس يوم {dp}")
            if data.get('warnings'):
                st.warning(f"⚠️ بدون مذكرة: **{', '.join(set(data['warnings']))}**")
            morning = [s for s in sessions if s['الفترة']=='صباحية']
            evening = [s for s in sessions if s['الفترة']=='مسائية']
            for pn,pl in [("☀️ الصباحية",morning),("🌙 المسائية",evening)]:
                if not pl: continue
                st.markdown(f"### {pn}")
                for s in pl:
                    typ = s['نوع']
                    css = {'روتيني':'slot-routine','تعليمي':'slot-teach','ناقص':'slot-warn'}.get(typ,'slot-routine')
                    ic = {'روتيني':'⏰','تعليمي':'📖','ناقص':'⚠️'}.get(typ,'⏰')
                    dom = s.get('المجال','—')
                    bdg = domain_badge(dom) if dom!='—' else ''
                    ext = ''
                    if typ=='تعليمي':
                        sn = s.get('رقم_الحصة','')
                        snb = f" <small style='color:#1565c0'>[{sn}]</small>" if sn else ''
                        ext = f"<br><small>📝 {s['الموضوع']}</small><br><small>🎯 {s['الكفاءة']}</small>{snb}"
                    elif typ=='ناقص':
                        ext = "<br><small>⚠ لا توجد مذكرة</small>"
                    st.markdown(f"""<div class="slot {css}">
                        <span style="font-size:1.3rem">{ic}</span>
                        <div style="flex:1"><strong>{s['النشاط']}</strong>
                        <span style="color:#888">({s['المدة']})</span>
                        {bdg}{ext}</div></div>""", unsafe_allow_html=True)


# ──── تبويب 4: التوزيع الذكي ────

with tab4:
    report = st.session_state.dist_report
    if not report:
        st.info("🧠 استخرج الدروس أولاً")
    else:
        st.markdown("### 📊 خريطة التوزيع الأسبوعي")
        st.markdown("> كيف تم توزيع الدروس المستخرجة على حصص الأسبوع")

        colors_list = ["#1565C0","#C62828","#2E7D32","#F57F17","#6A1B9A","#00838F","#E65100"]

        for subj, info in report.items():
            domain = get_domain_for(subj)
            dom_color = DOMAIN_COLORS.get(domain, "#666")
            status = info['حالة']
            req = info['مطلوب']
            avail = info['متوفر']
            dist = info['توزيع']

            st.markdown(f"**{status} {subj}** — {avail} درس ← {req} حصة — {domain_badge(domain)}",
                        unsafe_allow_html=True)

            if dist:
                # تصور مرئي للتوزيع
                slots_html = ""
                unique_lessons = list(dict.fromkeys(dist))
                for slot_label in dist:
                    idx = unique_lessons.index(slot_label)
                    col = colors_list[idx % len(colors_list)]
                    slots_html += (f'<div class="dist-slot" style="background:{col}">'
                                   f'{slot_label.replace("درس ","")}</div>')

                st.markdown(f'<div class="dist-row">{slots_html}'
                            f'<small style="color:#888;margin-right:8px;">'
                            f'{"→".join(unique_lessons)}</small></div>',
                            unsafe_allow_html=True)

            st.markdown("")

        # ملخص
        st.markdown("---")
        covered = sum(1 for v in report.values() if v['حالة']=='✅')
        total = len(report)
        st.markdown(f"""
        <div class="card card-{'green' if covered==total else 'amber'}">
            <h4>التغطية</h4>
            <div class="num">{covered}/{total}</div>
            <small>{'🎉 تغطية كاملة!' if covered==total else '⚠ بعض المواد ناقصة'}</small>
        </div>""", unsafe_allow_html=True)


# ──── تبويب 5: المجالات ────

with tab5:
    st.markdown("### 🗺️ المجالات التعليمية")
    domains = {}
    for subj,dom in DOMAIN_MAPPING.items():
        domains.setdefault(dom,[]).append(subj)
    cols5 = st.columns(2)
    for i,(dom,subjs) in enumerate(domains.items()):
        col = DOMAIN_COLORS.get(dom,"#666")
        total_h = sum(SUBJECT_WEEKLY_COUNT.get(s,0) for s in subjs)
        with cols5[i%2]:
            st.markdown(
                f'<div style="border:2px solid {col};border-radius:12px;'
                f'padding:1rem;margin:.5rem 0;">'
                f'<h4 style="color:{col};text-align:center;">'
                f'{dom} ({total_h} ح/أسبوع)</h4>', unsafe_allow_html=True)
            for s in subjs:
                cnt = SUBJECT_WEEKLY_COUNT.get(s,0)
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:8px;'
                    f'margin:4px 0;direction:rtl;">'
                    f'<span style="min-width:140px">{s}</span>'
                    f'<div style="background:{col}44;border-radius:4px;'
                    f'height:20px;width:{cnt*14}px;"></div>'
                    f'<span style="color:{col};font-weight:700;">{cnt}</span></div>',
                    unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")
    grand = sum(SUBJECT_WEEKLY_COUNT.values())
    st.markdown(f'<div class="card card-blue"><h4>المجموع</h4>'
                f'<div class="num">{grand} حصة</div></div>',
                unsafe_allow_html=True)


# ──── تبويب 6: محتوى الملف ────

with tab6:
    st.markdown("### 📄 المحتوى الخام للملف")
    txt = st.session_state.file_text
    if not txt:
        st.info("👆 ارفع ملفاً لرؤية محتواه")
    else:
        st.markdown(f"**طول النص:** {len(txt)} حرف | **الأسطر:** {txt.count(chr(10))+1}")
        st.text_area("المحتوى الكامل", txt, height=500, key="raw_content")

        # عرض رد AI إذا موجود
        if st.session_state.ai_raw:
            st.markdown("---")
            st.markdown("### 🧠 رد الذكاء الاصطناعي (الخام)")
            st.code(st.session_state.ai_raw, language="json")
