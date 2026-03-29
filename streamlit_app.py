# -*- coding: utf-8 -*-
"""🎓 الكراس اليومي v9 — Groq + HuggingFace + PDF/صور/Word"""

import streamlit as st
import re
import copy
import json
import base64
import requests
from io import BytesIO
from PIL import Image

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_OK = True
except Exception:
    DOCX_OK = False

try:
    import fitz  # pymupdf
    PDF_OK = True
except Exception:
    PDF_OK = False

# ╔══════════════════════════════════════════════════════╗
# ║ APIs — كل محرك بمفتاحه الخاص                         ║
# ╚══════════════════════════════════════════════════════╝

GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_TEXT_MODELS = ["llama-3.3-70b-versatile", "mixtral-8x7b-32768", "gemma2-9b-it"]
GROQ_VISION_MODELS = ["llama-3.2-90b-vision-preview", "llama-3.2-11b-vision-preview"]

HF_MODELS = [
    "mistralai/Mistral-7B-Instruct-v0.3",
    "meta-llama/Meta-Llama-3-8B-Instruct",
    "google/gemma-2-2b-it",
]


def groq_text_call(api_key, model, system_msg, user_msg):
    """استدعاء Groq للنصوص"""
    try:
        r = requests.post(
            GROQ_URL,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": model,
                "messages": [
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": user_msg},
                ],
                "temperature": 0.1,
                "max_tokens": 4000,
            },
            timeout=120,
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"], None
    except Exception as e:
        err_msg = ""
        if hasattr(e, 'response') and e.response is not None:
            err_msg = e.response.text[:300]
        else:
            err_msg = str(e)[:300]
        return None, f"Groq: {err_msg}"


def groq_vision_call(api_key, model, prompt, image_b64):
    """استدعاء Groq Vision للصور"""
    try:
        r = requests.post(
            GROQ_URL,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": model,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {
                            "url": f"data:image/jpeg;base64,{image_b64}"
                        }}
                    ]
                }],
                "temperature": 0.1,
                "max_tokens": 4000,
            },
            timeout=120,
        )
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"], None
    except Exception as e:
        err_msg = ""
        if hasattr(e, 'response') and e.response is not None:
            err_msg = e.response.text[:300]
        else:
            err_msg = str(e)[:300]
        return None, f"Groq Vision: {err_msg}"


def hf_text_call(token, model, prompt):
    """استدعاء Hugging Face للنصوص"""
    url = f"https://api-inference.huggingface.co/models/{model}"
    try:
        r = requests.post(
            url,
            headers={"Authorization": f"Bearer {token}"},
            json={
                "inputs": prompt,
                "parameters": {"max_new_tokens": 3000, "temperature": 0.1, "return_full_text": False},
            },
            timeout=120,
        )
        r.raise_for_status()
        result = r.json()
        if isinstance(result, list) and result:
            return result[0].get("generated_text", ""), None
        return str(result), None
    except Exception as e:
        err_msg = ""
        if hasattr(e, 'response') and e.response is not None:
            body = e.response.text[:300]
            if "loading" in body.lower() or "is currently loading" in body.lower():
                return None, "⏳ النموذج يتم تحميله... أعد المحاولة بعد 30 ثانية"
            err_msg = body
        else:
            err_msg = str(e)[:300]
        return None, f"HuggingFace: {err_msg}"


# ╔══════════════════════════════════════════════════════╗
# ║ قراءة الملفات — Word + PDF نصي + PDF صور + صور      ║
# ╚══════════════════════════════════════════════════════╝

def compress_image(data, max_size=1200, quality=80):
    img = Image.open(BytesIO(data))
    if img.mode == 'RGBA':
        img = img.convert('RGB')
    w, h = img.size
    if max(w, h) > max_size:
        ratio = max_size / max(w, h)
        img = img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
    buf = BytesIO()
    img.save(buf, format='JPEG', quality=quality)
    return buf.getvalue()


def to_base64(data):
    return base64.b64encode(compress_image(data)).decode()


def read_docx_text(file_bytes):
    if not DOCX_OK:
        return ""
    doc = Document(BytesIO(file_bytes))
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def read_pdf_text(file_bytes):
    """محاولة استخراج نص من PDF"""
    if not PDF_OK:
        return ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        texts = []
        for page in doc:
            t = page.get_text().strip()
            if t and len(t) > 20:
                texts.append(t)
        doc.close()
        return "\n".join(texts)
    except Exception:
        return ""


def pdf_to_images(file_bytes):
    """تحويل صفحات PDF إلى صور (للمسحوب ضوئياً)"""
    if not PDF_OK:
        return []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        images = []
        for i, page in enumerate(doc):
            # دقة عالية للقراءة
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("jpeg")
            images.append({
                "page": i + 1,
                "bytes": img_bytes,
                "b64": to_base64(img_bytes),
            })
        doc.close()
        return images
    except Exception:
        return []


# ╔══════════════════════════════════════════════════════╗
# ║ استخراج الدروس بالذكاء الاصطناعي                    ║
# ╚══════════════════════════════════════════════════════╝

def build_extraction_prompt(level_name, subjects_dict):
    subjects_list = "\n".join([f"- {s}" for s in subjects_dict.keys()])
    return f"""أنت محلل وثائق تربوية جزائرية متخصص. المستوى: {level_name}

استخرج كل الدروس الموجودة. المواد المتوقعة (استخدم هذه الأسماء بالضبط):
{subjects_list}

لكل درس استخرج:
1. "مادة": اسم المادة بالضبط من القائمة أعلاه
2. "موضوع": عنوان/موضوع الدرس
3. "كفاءة": مؤشر الكفاءة المستهدفة (إن لم يوجد اكتب "—")

مهم جداً:
- إذا وجدت مادة بأسماء مختلفة، وحّدها حسب القائمة
- استخرج كل الدروس حتى لو تكررت المادة
- أعد النتيجة كـ JSON فقط بدون أي نص إضافي

[{{"مادة":"...","موضوع":"...","كفاءة":"..."}}]"""


def parse_ai_json(raw_text):
    if not raw_text:
        return None
    for parser in [
        lambda: json.loads(raw_text),
        lambda: json.loads(re.search(r'```(?:json)?\s*([\s\S]*?)```', raw_text).group(1)),
        lambda: json.loads(re.search(r'\[[\s\S]*?\]', raw_text).group(0)),
    ]:
        try:
            result = parser()
            if isinstance(result, list):
                return result
        except Exception:
            pass
    return None


def ai_results_to_db(parsed_list):
    db = {}
    if not parsed_list:
        return db
    for item in parsed_list:
        if not isinstance(item, dict):
            continue
        subject = item.get("مادة", "").strip()
        topic = item.get("موضوع", "—").strip()
        indicator = item.get("كفاءة", "—").strip()
        if subject and topic:
            db.setdefault(subject, []).append({"موضوع": topic, "كفاءة": indicator})
    return db


def extract_from_text_groq(text, api_key, model, level, subjects):
    prompt = build_extraction_prompt(level, subjects)
    full_prompt = prompt + "\n\nالمحتوى:\n---\n" + text[:12000] + "\n---"
    raw, err = groq_text_call(api_key, model, "محلل وثائق تربوية. أجب بـ JSON فقط.", full_prompt)
    if err:
        return None, raw, err
    parsed = parse_ai_json(raw)
    if not parsed:
        return None, raw, "فشل تحليل رد AI كـ JSON"
    return ai_results_to_db(parsed), raw, None


def extract_from_text_hf(text, token, model, level, subjects):
    prompt = build_extraction_prompt(level, subjects)
    full_prompt = prompt + "\n\nالمحتوى:\n---\n" + text[:8000] + "\n---"
    raw, err = hf_text_call(token, model, full_prompt)
    if err:
        return None, raw, err
    parsed = parse_ai_json(raw)
    if not parsed:
        return None, raw, "فشل تحليل رد AI كـ JSON"
    return ai_results_to_db(parsed), raw, None


def extract_from_images_groq(images_b64, api_key, model, level, subjects, progress_cb=None):
    prompt = build_extraction_prompt(level, subjects)
    prompt += "\n\nاقرأ كل النص العربي من الصورة واستخرج الدروس."
    all_results = []
    all_raws = []
    for i, b64 in enumerate(images_b64):
        if progress_cb:
            progress_cb(i, len(images_b64))
        raw, err = groq_vision_call(api_key, model, prompt + f"\n(صورة {i+1}/{len(images_b64)})", b64)
        if err:
            all_raws.append(f"--- صفحة {i+1}: خطأ ---\n{err}")
            continue
        all_raws.append(f"--- صفحة {i+1} ---\n{raw}")
        parsed = parse_ai_json(raw)
        if parsed:
            all_results.extend(parsed)
    combined_raw = "\n\n".join(all_raws)
    if not all_results:
        return None, combined_raw, "لم يتم استخراج أي درس من الصور"
    return ai_results_to_db(all_results), combined_raw, None


def smart_process_file(file_bytes, filename, groq_key, hf_key, groq_tmodel, groq_vmodel, hf_model, level, subjects, progress_bar=None, status_text=None):
    """
    المعالج الذكي:
    - Word → نص → Groq نصي أو HF نصي
    - PDF نصي → نص → Groq نصي أو HF نصي
    - PDF صور → تحويل لصور → Groq Vision
    - صورة → Groq Vision مباشرة
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    result = {"db": None, "raw": None, "err": None, "method": ""}

    # ═══ تحديد أي مفتاح متاح ═══
    has_groq = bool(groq_key)
    has_hf = bool(hf_key)

    if not has_groq and not has_hf:
        result["err"] = "أدخل مفتاح API واحد على الأقل (Groq أو HuggingFace)"
        return result

    # ═══ Word ═══
    if ext == "docx":
        if status_text:
            status_text.text("📄 جارٍ قراءة ملف Word...")
        text = read_docx_text(file_bytes)
        if not text:
            result["err"] = "ملف Word فارغ أو تالف"
            return result
        if has_groq:
            result["method"] = "📄 Word → 🧠 Groq"
            if status_text:
                status_text.text("🧠 Groq يحلل النص...")
            db, raw, err = extract_from_text_groq(text, groq_key, groq_tmodel, level, subjects)
        else:
            result["method"] = "📄 Word → 🤗 HuggingFace"
            if status_text:
                status_text.text("🤗 HuggingFace يحلل النص...")
            db, raw, err = extract_from_text_hf(text, hf_key, hf_model, level, subjects)
        result["db"] = db
        result["raw"] = raw
        result["err"] = err

    # ═══ PDF ═══
    elif ext == "pdf":
        if status_text:
            status_text.text("📕 جارٍ قراءة PDF...")
        # محاولة 1: استخراج نص
        text = read_pdf_text(file_bytes)
        if text and len(text) > 100:
            # PDF نصي — نعالجه كنص
            if has_groq:
                result["method"] = "📕 PDF نصي → 🧠 Groq"
                if status_text:
                    status_text.text("🧠 Groq يحلل النص...")
                db, raw, err = extract_from_text_groq(text, groq_key, groq_tmodel, level, subjects)
            else:
                result["method"] = "📕 PDF نصي → 🤗 HuggingFace"
                if status_text:
                    status_text.text("🤗 HuggingFace يحلل النص...")
                db, raw, err = extract_from_text_hf(text, hf_key, hf_model, level, subjects)
            result["db"] = db
            result["raw"] = raw
            result["err"] = err
        else:
            # PDF مسحوب ضوئياً — نحوله لصور
            if not has_groq:
                result["err"] = (
                    "📕 هذا PDF مسحوب ضوئياً (صور). "
                    "يحتاج مفتاح Groq API لقراءته بالذكاء البصري.\n\n"
                    "أدخل مفتاح Groq أو صوّر الصفحات وارفعها كصور."
                )
                return result
            if not PDF_OK:
                result["err"] = "مكتبة pymupdf غير متوفرة لتحويل PDF"
                return result
            if status_text:
                status_text.text("📸 تحويل صفحات PDF إلى صور...")
            images = pdf_to_images(file_bytes)
            if not images:
                result["err"] = "فشل تحويل صفحات PDF إلى صور"
                return result
            result["method"] = f"📕 PDF ({len(images)} صفحة) → 📸 → 👁️ Groq Vision"
            if status_text:
                status_text.text(f"👁️ Groq Vision يقرأ {len(images)} صفحة...")
            b64_list = [img["b64"] for img in images]

            def on_progress(i, total):
                if progress_bar:
                    progress_bar.progress((i + 1) / total)
                if status_text:
                    status_text.text(f"👁️ صفحة {i+1}/{total}...")

            db, raw, err = extract_from_images_groq(
                b64_list, groq_key, groq_vmodel, level, subjects, on_progress
            )
            result["db"] = db
            result["raw"] = raw
            result["err"] = err

    # ═══ صورة ═══
    elif ext in ("jpg", "jpeg", "png", "bmp", "webp", "tiff"):
        if not has_groq:
            result["err"] = "قراءة الصور تحتاج مفتاح Groq API (يدعم Vision)"
            return result
        result["method"] = "🖼️ صورة → 👁️ Groq Vision"
        if status_text:
            status_text.text("👁️ Groq Vision يقرأ الصورة...")
        b64 = to_base64(file_bytes)
        db, raw, err = extract_from_images_groq(
            [b64], groq_key, groq_vmodel, level, subjects
        )
        result["db"] = db
        result["raw"] = raw
        result["err"] = err

    else:
        result["err"] = f"صيغة غير مدعومة: .{ext}"

    return result


# ╔══════════════════════════════════════════════════════╗
# ║ المستويات الدراسية                                   ║
# ╚══════════════════════════════════════════════════════╝

LEVELS = {
    "قسم التحضيري": {
        "رت": ["الاستقبال", "الاستراحة", "تهيئة الخروج", "نهاية الخروج"],
        "مج": {
            "المجال اللغوي": "#1565C0",
            "المجال الرياضي": "#C62828",
            "المجال العلمي": "#2E7D32",
            "المجال الاجتماعي": "#F57F17",
            "المجال الفني": "#6A1B9A",
            "المجال البدني والإيقاعي": "#00838F"
        },
        "مو": {
            "تعبير شفوي": {"ج": "المجال اللغوي", "ح": 3},
            "مبادئ القراءة": {"ج": "المجال اللغوي", "ح": 4},
            "تخطيط": {"ج": "المجال اللغوي", "ح": 2},
            "رياضيات": {"ج": "المجال الرياضي", "ح": 5},
            "ت علمية وتكنولوجية": {"ج": "المجال العلمي", "ح": 4},
            "ت إسلامية": {"ج": "المجال الاجتماعي", "ح": 2},
            "ت مدنية": {"ج": "المجال الاجتماعي", "ح": 2},
            "تربية تشكيلية": {"ج": "المجال الفني", "ح": 2},
            "موسيقى وإنشاد": {"ج": "المجال الفني", "ح": 2},
            "مسرح وعرائس": {"ج": "المجال الفني", "ح": 2},
            "ت بدنية": {"ج": "المجال البدني والإيقاعي", "ح": 4},
            "ت إيقاعية": {"ج": "المجال البدني والإيقاعي", "ح": 2},
        },
        "تو": {
            "الأحد": [
                {"ن": "الاستقبال", "م": "15 د", "ف": "ص"},
                {"ن": "تعبير شفوي", "م": "30 د", "ف": "ص"},
                {"ن": "مبادئ القراءة", "م": "30 د", "ف": "ص"},
                {"ن": "رياضيات", "م": "30 د", "ف": "ص"},
                {"ن": "ت علمية وتكنولوجية", "م": "30 د", "ف": "ص"},
                {"ن": "ت إسلامية", "م": "30 د", "ف": "ص"},
                {"ن": "تهيئة الخروج", "م": "15 د", "ف": "ص"},
                {"ن": "الاستقبال", "م": "15 د", "ف": "م"},
                {"ن": "مسرح وعرائس", "م": "30 د", "ف": "م"},
                {"ن": "تربية تشكيلية", "م": "30 د", "ف": "م"},
                {"ن": "ت بدنية", "م": "30 د", "ف": "م"},
                {"ن": "نهاية الخروج", "م": "15 د", "ف": "م"},
            ],
            "الإثنين": [
                {"ن": "الاستقبال", "م": "15 د", "ف": "ص"},
                {"ن": "رياضيات", "م": "30 د", "ف": "ص"},
                {"ن": "تعبير شفوي", "م": "30 د", "ف": "ص"},
                {"ن": "تخطيط", "م": "30 د", "ف": "ص"},
                {"ن": "ت علمية وتكنولوجية", "م": "30 د", "ف": "ص"},
                {"ن": "ت مدنية", "م": "30 د", "ف": "ص"},
                {"ن": "تهيئة الخروج", "م": "15 د", "ف": "ص"},
                {"ن": "الاستقبال", "م": "15 د", "ف": "م"},
                {"ن": "مسرح وعرائس", "م": "30 د", "ف": "م"},
                {"ن": "تربية تشكيلية", "م": "30 د", "ف": "م"},
                {"ن": "ت بدنية", "م": "30 د", "ف": "م"},
                {"ن": "نهاية الخروج", "م": "15 د", "ف": "م"},
            ],
            "الثلاثاء": [
                {"ن": "الاستقبال", "م": "15 د", "ف": "ص"},
                {"ن": "تعبير شفوي", "م": "30 د", "ف": "ص"},
                {"ن": "مبادئ القراءة", "م": "30 د", "ف": "ص"},
                {"ن": "رياضيات", "م": "30 د", "ف": "ص"},
                {"ن": "ت إسلامية", "م": "30 د", "ف": "ص"},
                {"ن": "ت بدنية", "م": "30 د", "ف": "ص"},
                {"ن": "تهيئة الخروج", "م": "15 د", "ف": "ص"},
            ],
            "الأربعاء": [
                {"ن": "الاستقبال", "م": "15 د", "ف": "ص"},
                {"ن": "رياضيات", "م": "30 د", "ف": "ص"},
                {"ن": "مبادئ القراءة", "م": "30 د", "ف": "ص"},
                {"ن": "تخطيط", "م": "30 د", "ف": "ص"},
                {"ن": "ت علمية وتكنولوجية", "م": "30 د", "ف": "ص"},
                {"ن": "ت مدنية", "م": "30 د", "ف": "ص"},
                {"ن": "تهيئة الخروج", "م": "15 د", "ف": "ص"},
                {"ن": "الاستقبال", "م": "15 د", "ف": "م"},
                {"ن": "ت إيقاعية", "م": "30 د", "ف": "م"},
                {"ن": "موسيقى وإنشاد", "م": "30 د", "ف": "م"},
                {"ن": "ت بدنية", "م": "30 د", "ف": "م"},
                {"ن": "نهاية الخروج", "م": "15 د", "ف": "م"},
            ],
            "الخميس": [
                {"ن": "الاستقبال", "م": "15 د", "ف": "ص"},
                {"ن": "مبادئ القراءة", "م": "30 د", "ف": "ص"},
                {"ن": "رياضيات", "م": "30 د", "ف": "ص"},
                {"ن": "ت علمية وتكنولوجية", "م": "30 د", "ف": "ص"},
                {"ن": "ت إيقاعية", "م": "30 د", "ف": "ص"},
                {"ن": "موسيقى وإنشاد", "م": "30 د", "ف": "ص"},
                {"ن": "تهيئة الخروج", "م": "15 د", "ف": "ص"},
            ],
        },
    },
    "السنة 1 ابتدائي": {
        "رت": ["الاستقبال", "الاستراحة", "تهيئة الخروج", "نهاية الخروج"],
        "مج": {
            "اللغة العربية": "#1565C0",
            "الرياضيات": "#C62828",
            "التربية العلمية": "#2E7D32",
            "التربية الاجتماعية": "#F57F17",
            "التربية الفنية": "#6A1B9A",
            "التربية البدنية": "#00838F"
        },
        "مو": {
            "قراءة": {"ج": "اللغة العربية", "ح": 6},
            "تعبير شفوي": {"ج": "اللغة العربية", "ح": 2},
            "كتابة وخط": {"ج": "اللغة العربية", "ح": 3},
            "محفوظات": {"ج": "اللغة العربية", "ح": 1},
            "رياضيات": {"ج": "الرياضيات", "ح": 5},
            "تربية علمية": {"ج": "التربية العلمية", "ح": 2},
            "تربية إسلامية": {"ج": "التربية الاجتماعية", "ح": 2},
            "تربية مدنية": {"ج": "التربية الاجتماعية", "ح": 1},
            "تربية فنية": {"ج": "التربية الفنية", "ح": 2},
            "تربية موسيقية": {"ج": "التربية الفنية", "ح": 1},
            "تربية بدنية": {"ج": "التربية البدنية", "ح": 2},
        },
        "تو": {},
    },
    "السنة 2 ابتدائي": {
        "رت": ["الاستقبال", "الاستراحة", "تهيئة الخروج", "نهاية الخروج"],
        "مج": {
            "اللغة العربية": "#1565C0",
            "الرياضيات": "#C62828",
            "التربية العلمية": "#2E7D32",
            "التربية الاجتماعية": "#F57F17",
            "التربية الفنية": "#6A1B9A",
            "التربية البدنية": "#00838F"
        },
        "مو": {
            "قراءة": {"ج": "اللغة العربية", "ح": 5},
            "تعبير شفوي": {"ج": "اللغة العربية", "ح": 2},
            "كتابة وخط": {"ج": "اللغة العربية", "ح": 3},
            "إملاء": {"ج": "اللغة العربية", "ح": 1},
            "محفوظات": {"ج": "اللغة العربية", "ح": 1},
            "رياضيات": {"ج": "الرياضيات", "ح": 5},
            "تربية علمية": {"ج": "التربية العلمية", "ح": 2},
            "تربية إسلامية": {"ج": "التربية الاجتماعية", "ح": 2},
            "تربية مدنية": {"ج": "التربية الاجتماعية", "ح": 1},
            "تربية فنية": {"ج": "التربية الفنية", "ح": 2},
            "تربية موسيقية": {"ج": "التربية الفنية", "ح": 1},
            "تربية بدنية": {"ج": "التربية البدنية", "ح": 2},
        },
        "تو": {},
    },
    "السنة 3 ابتدائي": {
        "رت": ["الاستقبال", "الاستراحة", "تهيئة الخروج", "نهاية الخروج"],
        "مج": {
            "اللغة العربية": "#1565C0",
            "اللغة الفرنسية": "#1976D2",
            "الرياضيات": "#C62828",
            "التربية العلمية": "#2E7D32",
            "التربية الاجتماعية": "#F57F17",
            "التربية الفنية": "#6A1B9A",
            "التربية البدنية": "#00838F"
        },
        "مو": {
            "قراءة": {"ج": "اللغة العربية", "ح": 4},
            "تعبير شفوي": {"ج": "اللغة العربية", "ح": 2},
            "كتابة وإملاء": {"ج": "اللغة العربية", "ح": 2},
            "قواعد": {"ج": "اللغة العربية", "ح": 1},
            "محفوظات": {"ج": "اللغة العربية", "ح": 1},
            "فرنسية": {"ج": "اللغة الفرنسية", "ح": 3},
            "رياضيات": {"ج": "الرياضيات", "ح": 5},
            "تربية علمية": {"ج": "التربية العلمية", "ح": 2},
            "تربية إسلامية": {"ج": "التربية الاجتماعية", "ح": 1},
            "تربية مدنية": {"ج": "التربية الاجتماعية", "ح": 1},
            "تاريخ وجغرافيا": {"ج": "التربية الاجتماعية", "ح": 1},
            "تربية فنية": {"ج": "التربية الفنية", "ح": 1},
            "تربية موسيقية": {"ج": "التربية الفنية", "ح": 1},
            "تربية بدنية": {"ج": "التربية البدنية", "ح": 2},
        },
        "تو": {},
    },
    "السنة 4 ابتدائي": {
        "رت": ["الاستقبال", "الاستراحة", "تهيئة الخروج", "نهاية الخروج"],
        "مج": {
            "اللغة العربية": "#1565C0",
            "اللغة الفرنسية": "#1976D2",
            "الرياضيات": "#C62828",
            "التربية العلمية": "#2E7D32",
            "التربية الاجتماعية": "#F57F17",
            "التربية الفنية": "#6A1B9A",
            "التربية البدنية": "#00838F"
        },
        "مو": {
            "قراءة ودراسة نص": {"ج": "اللغة العربية", "ح": 3},
            "قواعد صرفية ونحوية": {"ج": "اللغة العربية", "ح": 2},
            "تعبير كتابي": {"ج": "اللغة العربية", "ح": 1},
            "تعبير شفوي": {"ج": "اللغة العربية", "ح": 1},
            "إملاء": {"ج": "اللغة العربية", "ح": 1},
            "محفوظات": {"ج": "اللغة العربية", "ح": 1},
            "فرنسية": {"ج": "اللغة الفرنسية", "ح": 3},
            "رياضيات": {"ج": "الرياضيات", "ح": 5},
            "تربية علمية": {"ج": "التربية العلمية", "ح": 2},
            "تربية إسلامية": {"ج": "التربية الاجتماعية", "ح": 1},
            "تربية مدنية": {"ج": "التربية الاجتماعية", "ح": 1},
            "تاريخ": {"ج": "التربية الاجتماعية", "ح": 1},
            "جغرافيا": {"ج": "التربية الاجتماعية", "ح": 1},
            "تربية فنية": {"ج": "التربية الفنية", "ح": 1},
            "تربية موسيقية": {"ج": "التربية الفنية", "ح": 1},
            "تربية بدنية": {"ج": "التربية البدنية", "ح": 2},
        },
        "تو": {},
    },
    "السنة 5 ابتدائي": {
        "رت": ["الاستقبال", "الاستراحة", "تهيئة الخروج", "نهاية الخروج"],
        "مج": {
            "اللغة العربية": "#1565C0",
            "اللغة الفرنسية": "#1976D2",
            "الرياضيات": "#C62828",
            "التربية العلمية": "#2E7D32",
            "التربية الاجتماعية": "#F57F17",
            "التربية الفنية": "#6A1B9A",
            "التربية البدنية": "#00838F"
        },
        "مو": {
            "قراءة ودراسة نص": {"ج": "اللغة العربية", "ح": 3},
            "قواعد صرفية ونحوية": {"ج": "اللغة العربية", "ح": 2},
            "تعبير كتابي": {"ج": "اللغة العربية", "ح": 1},
            "تعبير شفوي": {"ج": "اللغة العربية", "ح": 1},
            "إملاء": {"ج": "اللغة العربية", "ح": 1},
            "محفوظات": {"ج": "اللغة العربية", "ح": 1},
            "فرنسية": {"ج": "اللغة الفرنسية", "ح": 3},
            "رياضيات": {"ج": "الرياضيات", "ح": 5},
            "تربية علمية": {"ج": "التربية العلمية", "ح": 2},
            "تربية إسلامية": {"ج": "التربية الاجتماعية", "ح": 1},
            "تربية مدنية": {"ج": "التربية الاجتماعية", "ح": 1},
            "تاريخ": {"ج": "التربية الاجتماعية", "ح": 1},
            "جغرافيا": {"ج": "التربية الاجتماعية", "ح": 1},
            "تربية فنية": {"ج": "التربية الفنية", "ح": 1},
            "تربية موسيقية": {"ج": "التربية الفنية", "ح": 1},
            "تربية بدنية": {"ج": "التربية البدنية", "ح": 2},
        },
        "تو": {},
    },
}


# ╔══════════════════════════════════════════════════════╗
# ║ دوال المستوى والتوقيت                                ║
# ╚══════════════════════════════════════════════════════╝

DAYS = ["الأحد", "الإثنين", "الثلاثاء", "الأربعاء", "الخميس"]
HALF_DAYS = {"الثلاثاء", "الخميس"}


def get_cfg():
    lv = st.session_state.get('lv', 'قسم التحضيري')
    return st.session_state.get('cfgs', LEVELS).get(lv, LEVELS['قسم التحضيري'])


def auto_schedule(subjects, routine):
    ss = []
    for s, i in subjects.items():
        ss.extend([s] * i["ح"])
    sch = {}
    idx = 0
    for day in DAYS:
        plan = []
        is_full = day not in HALF_DAYS
        plan.append({"ن": routine[0] if routine else "الاستقبال", "م": "15 د", "ف": "ص"})
        for _ in range(5):
            if idx < len(ss):
                plan.append({"ن": ss[idx], "م": "45 د", "ف": "ص"})
                idx += 1
        plan.append({"ن": "تهيئة الخروج", "م": "15 د", "ف": "ص"})
        if is_full:
            plan.append({"ن": routine[0] if routine else "الاستقبال", "م": "15 د", "ف": "م"})
            for _ in range(3):
                if idx < len(ss):
                    plan.append({"ن": ss[idx], "م": "45 د", "ف": "م"})
                    idx += 1
            plan.append({"ن": "نهاية الخروج", "م": "15 د", "ف": "م"})
        sch[day] = plan
    return sch


def get_schedule():
    c = get_cfg()
    s = c.get("تو", {})
    return s if s else auto_schedule(c["مو"], c["رت"])


def get_subjects():
    return get_cfg().get("مو", {})


def get_domains():
    return get_cfg().get("مج", {})


def get_routine():
    return get_cfg().get("رت", [])


def domain_for(activity):
    m = get_subjects()
    return m[activity]["ج"] if activity in m else "—"


def domain_color(domain):
    return get_domains().get(domain, "#666")


def domain_badge(domain):
    c = domain_color(domain)
    return (f'<span style="display:inline-block;padding:2px 8px;border-radius:10px;'
            f'font-size:.7rem;font-weight:600;background:{c}22;color:{c};'
            f'border:1px solid {c}33;">{domain}</span>')


# ╔══════════════════════════════════════════════════════╗
# ║ التوزيع الذكي للحصص                                 ║
# ╚══════════════════════════════════════════════════════╝

def distribute_lessons(raw_db):
    subjects = get_subjects()
    distributed = {}
    report = {}
    for subj, info in subjects.items():
        required = info["ح"]
        available = raw_db.get(subj, [])
        count = len(available)
        if count == 0:
            distributed[subj] = []
            report[subj] = {"مطلوب": required, "متوفر": 0, "حالة": False}
            continue
        result = []
        if count >= required:
            result = [l.copy() for l in available[:required]]
        else:
            per = required / count
            for i, lesson in enumerate(available):
                slots = round((i + 1) * per) - round(i * per)
                for _ in range(slots):
                    result.append(lesson.copy())
        distributed[subj] = result
        report[subj] = {"مطلوب": required, "متوفر": count, "حالة": bool(result)}
    return distributed, report


# ╔══════════════════════════════════════════════════════╗
# ║ بناء القالب وحقن البيانات                           ║
# ╚══════════════════════════════════════════════════════╝

def _rtl(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn('w:bidi'), {}))


def _cell_write(cell, text, bold=False, size=10, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Sakkal Majalla"
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rPr.append(rPr.makeelement(qn('w:rFonts'), {qn('w:cs'): 'Sakkal Majalla'}))


def _shade_row(row, hex_color):
    for c in row.cells:
        tc = c._tc.get_or_add_tcPr()
        tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): hex_color, qn('w:val'): 'clear'}))


def _add_period_table(doc, title, start_idx, count):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl(h)
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 51, 102)

    headers = ['مؤشرات الكفاءة', 'عنوان الدرس', 'الميدان', 'النشاط', 'المدة']
    tbl = doc.add_table(rows=1 + count, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Cm(5.5), Cm(5), Cm(3.5), Cm(3.5), Cm(2)]
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    hdr = tbl.rows[0]
    _shade_row(hdr, "1F4E79")
    for i, t in enumerate(headers):
        _cell_write(hdr.cells[i], t, True, 11, RGBColor(255, 255, 255))

    for j in range(count):
        n = start_idx + j
        data_row = tbl.rows[1 + j]
        if j % 2 == 0:
            _shade_row(data_row, "EDF2F9")
        placeholders = [f'{{{{k_{n}}}}}', f'{{{{t_{n}}}}}', f'{{{{d_{n}}}}}', f'{{{{n_{n}}}}}', f'{{{{m_{n}}}}}']
        for i, ph in enumerate(placeholders):
            _cell_write(data_row.cells[i], ph, size=9)


def create_template(day_name=None):
    if not DOCX_OK:
        return None
    level = st.session_state.get('lv', '')
    schedule = get_schedule()
    morning_count = 7
    evening_count = 5
    if day_name and day_name in schedule:
        plan = schedule[day_name]
        morning_count = sum(1 for s in plan if s['ف'] == 'ص')
        evening_count = sum(1 for s in plan if s['ف'] == 'م')

    doc = Document()
    for sec in doc.sections:
        sec._sectPr.append(sec._sectPr.makeelement(qn('w:bidi'), {}))

    for text, sz, bold in [('الجمهورية الجزائرية الديمقراطية الشعبية', 12, True),
                           ('وزارة التربية الوطنية', 11, False)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _rtl(p)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(sz)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl(title_p)
    title_run = title_p.add_run(f'الكراس اليومي — {level}')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    info_tbl = doc.add_table(rows=1, cols=3)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _cell_write(info_tbl.rows[0].cells[2], 'اليوم: {{day}}', True, 12)
    _cell_write(info_tbl.rows[0].cells[1], 'التاريخ: {{date}}', size=11)
    _cell_write(info_tbl.rows[0].cells[0], 'الأسبوع: {{week}}', size=11)

    doc.add_paragraph('')
    _add_period_table(doc, '☀ الفترة الصباحية', 1, morning_count)
    if evening_count > 0:
        doc.add_paragraph('')
        _add_period_table(doc, '🌙 الفترة المسائية', morning_count + 1, evening_count)

    doc.add_paragraph('')
    notes_p = doc.add_paragraph()
    notes_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _rtl(notes_p)
    notes_p.add_run('ملاحظات: ' + '.' * 80).font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _safe_replace(paragraph, old_text, new_text):
    if old_text not in paragraph.text:
        return
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return
    full = paragraph.text.replace(old_text, new_text)
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = full


def build_daily_planner(day, template_bytes, distributed_db, week_num="", date_str=""):
    if not DOCX_OK:
        return None, [], []
    schedule = get_schedule()
    plan = schedule.get(day, [])
    if not plan:
        return None, [], []
    routine = get_routine()
    doc = Document(BytesIO(template_bytes))
    replacements = {"{{day}}": day, "{{date}}": date_str, "{{week}}": week_num}
    sessions_info = []
    warnings = []

    for i, session in enumerate(plan, 1):
        activity = session["ن"]
        duration = session["م"]
        period = "صباحية" if session["ف"] == "ص" else "مسائية"
        domain = domain_for(activity)

        replacements[f"{{{{m_{i}}}}}"] = duration
        replacements[f"{{{{n_{i}}}}}"] = activity

        info = {
            "رقم": i,
            "النشاط": activity,
            "المدة": duration,
            "الفترة": period,
            "المجال": domain,
            "نوع": "روتيني",
            "الموضوع": "—",
            "الكفاءة": "—"
        }

        if activity in routine:
            replacements[f"{{{{t_{i}}}}}"] = "—"
            replacements[f"{{{{k_{i}}}}}"] = "—"
            replacements[f"{{{{d_{i}}}}}"] = "—"
        elif activity in distributed_db and distributed_db[activity]:
            lesson = distributed_db[activity].pop(0)
            topic = lesson.get('موضوع', '—')
            indicator = lesson.get('كفاءة', '—')
            replacements[f"{{{{t_{i}}}}}"] = topic
            replacements[f"{{{{k_{i}}}}}"] = indicator
            replacements[f"{{{{d_{i}}}}}"] = domain
            info.update({"نوع": "تعليمي", "الموضوع": topic, "الكفاءة": indicator})
        else:
            replacements[f"{{{{t_{i}}}}}"] = "⚠ لا توجد مذكرة"
            replacements[f"{{{{k_{i}}}}}"] = "⚠"
            replacements[f"{{{{d_{i}}}}}"] = domain
            info["نوع"] = "ناقص"
            warnings.append(activity)

        sessions_info.append(info)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in para.text:
                            _safe_replace(para, key, str(value))

    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                _safe_replace(para, key, str(value))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), sessions_info, warnings


# ╔══════════════════════════════════════════════════════╗
# ║ واجهة Streamlit                                      ║
# ╚══════════════════════════════════════════════════════╝

st.set_page_config(page_title="الكراس اليومي 🎓", page_icon="🎓", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
.main .block-container{direction:rtl;text-align:right}
h1,h2,h3{text-align:center!important}
.card{padding:1rem;border-radius:12px;text-align:center;margin:.4rem 0;box-shadow:0 2px 8px rgba(0,0,0,.08)}
.card h4{margin:0 0 .3rem 0;font-size:.9rem}
.card .num{font-size:2rem;font-weight:700}
.card-blue{background:linear-gradient(135deg,#1F4E79,#2E75B6);color:#fff}
.card-green{background:linear-gradient(135deg,#2E7D32,#43A047);color:#fff}
.card-amber{background:linear-gradient(135deg,#E65100,#FF9800);color:#fff}
.card-purple{background:linear-gradient(135deg,#4A148C,#7B1FA2);color:#fff}
.slot{display:flex;align-items:center;gap:.8rem;padding:.6rem 1rem;margin:.2rem 0;border-radius:8px;direction:rtl}
.slot-r{background:#f5f5f5;border-right:4px solid #9e9e9e}
.slot-t{background:#e3f2fd;border-right:4px solid #1565c0}
.slot-w{background:#fff3e0;border-right:4px solid #e65100}
.stDownloadButton>button{width:100%;background:linear-gradient(135deg,#1F4E79,#2E75B6)!important;color:#fff!important;border:none!important;border-radius:8px!important}
[data-testid="stSidebar"]{direction:rtl;text-align:right}
footer{visibility:hidden}
.ok-box{background:#e8f5e9;border:1px solid #4caf50;border-radius:10px;padding:1rem;text-align:center}
.api-box{border:1px solid #ddd;border-radius:10px;padding:1rem;margin:.5rem 0}
</style>
""", unsafe_allow_html=True)

# Session State
for key in ['db', 'gen', 'raw', 'lv', 'cfgs', 'method']:
    if key not in st.session_state:
        if key == 'gen':
            st.session_state[key] = {}
        elif key == 'lv':
            st.session_state[key] = 'قسم التحضيري'
        elif key == 'cfgs':
            st.session_state[key] = copy.deepcopy(LEVELS)
        else:
            st.session_state[key] = None

# ════════════════ الشريط الجانبي ════════════════
with st.sidebar:
    st.markdown("## 🎓 المستوى الدراسي")
    level_names = list(LEVELS.keys())
    selected_level = st.selectbox(
        "اختر المستوى",
        level_names,
        index=level_names.index(st.session_state.lv)
    )
    if selected_level != st.session_state.lv:
        st.session_state.lv = selected_level
        st.session_state.db = None
        st.session_state.gen = {}
        st.rerun()

    cfg = get_cfg()
    total_sessions = sum(i["ح"] for i in cfg["مو"].values())
    st.info(f"📘 {len(cfg['مو'])} مادة | 📖 {total_sessions} حصة/أسبوع")

    st.markdown("---")
    week_num = st.text_input("📅 رقم الأسبوع", placeholder="مثال: 10")
    date_str = st.text_input("📆 التاريخ", placeholder="مثال: 2024/12/01")

    st.markdown("---")

    # ═══ مفاتيح API منفصلة ═══
    st.markdown("### 🔑 مفاتيح AI")
    st.markdown("**🟢 Groq** (نصوص + صور + PDF)")
    groq_key = st.text_input(
        "مفتاح Groq", type="password",
        help="احصل عليه مجاناً من console.groq.com",
        key="groq_key_input"
    )
    groq_text_model = st.selectbox("نموذج نصي", GROQ_TEXT_MODELS, key="gtm")
    groq_vision_model = st.selectbox("نموذج بصري", GROQ_VISION_MODELS, key="gvm")

    st.markdown("**🟡 HuggingFace** (نصوص فقط — بديل مجاني)")
    hf_key = st.text_input(
        "مفتاح HuggingFace", type="password",
        help="احصل عليه مجاناً من huggingface.co/settings/tokens",
        key="hf_key_input"
    )
    hf_model = st.selectbox("نموذج HF", HF_MODELS, key="hfm")

    st.markdown("---")

    # حالة المكتبات
    st.markdown("### 📦 حالة النظام")
    st.markdown(f"{'✅' if DOCX_OK else '❌'} Word (python-docx)")
    st.markdown(f"{'✅' if PDF_OK else '❌'} PDF (pymupdf)")
    st.markdown(f"{'✅' if groq_key else '⬜'} مفتاح Groq")
    st.markdown(f"{'✅' if hf_key else '⬜'} مفتاح HuggingFace")

    st.markdown("---")
    uploaded = st.file_uploader(
        "📤 رفع المذكرات",
        type=["docx", "pdf", "jpg", "jpeg", "png", "bmp", "webp"],
    )
    st.caption("🎓 v9.0")

# ════════════════ معالجة الرفع ════════════════
if uploaded:
    file_bytes = uploaded.read()
    if st.session_state.get('_last_file') != uploaded.name:
        st.session_state._file_bytes = file_bytes
        st.session_state._last_file = uploaded.name
        st.session_state.db = None
        st.session_state.gen = {}
        st.session_state.raw = None
        st.session_state.method = None

# ════════════════ العنوان الرئيسي ════════════════
current_level = st.session_state.lv
st.markdown(f"""
<h1 style="background:linear-gradient(135deg,#1F4E79,#2E75B6);
-webkit-background-clip:text;-webkit-text-fill-color:transparent;
font-size:2.2rem;">🎓 الكراس اليومي — {current_level}</h1>
<p style="text-align:center;color:#888;">
📄 Word • 📕 PDF • 🖼️ صور • 🟢 Groq • 🟡 HuggingFace
</p>
""", unsafe_allow_html=True)

# ════════════════ التبويبات ════════════════
tab1, tab2, tab3, tab4 = st.tabs(["🧠 استخراج الدروس", "📅 توليد الكراسات", "⚙️ إعداد المستوى", "🗺️ المجالات والتوقيت"])

# ──── تبويب 1: الاستخراج ────
with tab1:
    if not uploaded:
        st.info("👆 ارفع ملف المذكرات من الشريط الجانبي")
        st.markdown("""
        ### 📁 الصيغ المدعومة
        | الصيغة | 🟢 Groq | 🟡 HuggingFace |
        |--------|---------|----------------|
        | 📄 Word (.docx) | ✅ نصي | ✅ نصي |
        | 📕 PDF نصي | ✅ نصي | ✅ نصي |
        | 📕 PDF صور (مسحوب) | ✅ بصري | ❌ |
        | 🖼️ صورة (.jpg .png) | ✅ بصري | ❌ |

        ### 🔑 كيف تحصل على المفاتيح؟
        - **Groq** (مجاني): [console.groq.com](https://console.groq.com)
        - **HuggingFace** (مجاني): [huggingface.co/settings/tokens](https://huggingface.co/settings/tokens)
        """)
    else:
        file_ext = uploaded.name.lower().rsplit('.', 1)[-1]
        file_size = len(st.session_state._file_bytes) // 1024
        st.markdown(f"### 📁 {uploaded.name} ({file_size} KB)")

        # معاينة
        if file_ext in ('jpg', 'jpeg', 'png', 'bmp', 'webp'):
            st.image(st.session_state._file_bytes, use_container_width=True)
        elif file_ext == 'pdf' and PDF_OK:
            with st.expander("👁️ معاينة الصفحات"):
                preview_images = pdf_to_images(st.session_state._file_bytes)
                if preview_images:
                    cols = st.columns(min(len(preview_images), 3))
                    for idx, img_data in enumerate(preview_images[:6]):
                        with cols[idx % 3]:
                            st.image(img_data["bytes"], caption=f"صفحة {img_data['page']}")

        # التحقق من المفاتيح
        has_any_key = bool(groq_key) or bool(hf_key)
        if not has_any_key:
            st.warning("🔑 أدخل مفتاح API واحد على الأقل في الشريط الجانبي")

        # زر الاستخراج
        if st.button("🧠 استخراج الدروس تلقائياً", type="primary", use_container_width=True, disabled=not has_any_key):
            progress_bar = st.progress(0)
            status_msg = st.empty()
            result = smart_process_file(
                file_bytes=st.session_state._file_bytes,
                filename=uploaded.name,
                groq_key=groq_key,
                hf_key=hf_key,
                groq_tmodel=groq_text_model,
                groq_vmodel=groq_vision_model,
                hf_model=hf_model,
                level=current_level,
                subjects=get_subjects(),
                progress_bar=progress_bar,
                status_text=status_msg,
            )
            progress_bar.progress(1.0)
            status_msg.empty()
            st.session_state.raw = result["raw"]
            st.session_state.method = result["method"]
            if result["err"]:
                st.error(f"❌ {result['err']}")
            elif result["db"]:
                distributed, report = distribute_lessons(result["db"])
                st.session_state.db = distributed
                st.session_state.gen = {}
                total_extracted = sum(len(v) for v in result["db"].values())
                st.success(f"✅ تم استخراج {total_extracted} درس بنجاح!")
                st.rerun()
            else:
                st.error("❌ لم يتم العثور على دروس في الملف")

        # عرض النتائج
        db = st.session_state.db
        if db:
            if st.session_state.get('method'):
                st.caption(f"🔄 المسار: {st.session_state.method}")

            total = sum(len(v) for v in db.values())
            subjects_cfg = get_subjects()
            matched = set(k for k, v in db.items() if v) & set(subjects_cfg.keys())
            missing = set(subjects_cfg.keys()) - set(k for k, v in db.items() if v)

            col1, col2, col3 = st.columns(3)
            for col, title, num, css_class in [
                (col1, "📖 حصص", total, "card-green"),
                (col2, "✅ مغطاة", len(matched), "card-purple"),
                (col3, "⚠ ناقصة", len(missing), "card-amber"),
            ]:
                with col:
                    st.markdown(
                        f'<div class="card {css_class}"><h4>{title}</h4>'
                        f'<div class="num">{num}</div></div>',
                        unsafe_allow_html=True,
                    )

            if missing:
                st.warning(f"⚠️ مواد ناقصة: **{' ، '.join(missing)}**")

            for subj in sorted(db.keys()):
                lessons = db[subj]
                if not lessons:
                    continue
                domain = domain_for(subj)
                with st.expander(f"✅ {subj} — {len(lessons)} حصة — {domain}"):
                    for j, lesson in enumerate(lessons, 1):
                        st.markdown(
                            f"**{j}.** 📝 {lesson.get('موضوع', '—')}\n\n"
                            f"🎯 {lesson.get('كفاءة', '—')}"
                        )
                        if j < len(lessons):
                            st.divider()

            if st.session_state.raw:
                with st.expander("🧠 رد AI الخام"):
                    st.code(st.session_state.raw, language="json")

# ──── تبويب 2: التوليد ────
with tab2:
    db = st.session_state.db
    if not db or not any(db.values()):
        st.info("🧠 استخرج الدروس أولاً من التبويب الأول")
    elif not DOCX_OK:
        st.error("❌ مكتبة python-docx غير متوفرة — لا يمكن إنشاء ملفات Word")
    else:
        schedule = get_schedule()
        days = list(schedule.keys())
        st.markdown("### 📅 اختر الأيام")
        day_cols = st.columns(len(days))
        selected_days = []
        for i, day in enumerate(days):
            routine = get_routine()
            teaching_count = sum(1 for s in schedule[day] if s["ن"] not in routine)
            with day_cols[i]:
                if st.checkbox(f"{day} ({teaching_count})", key=f"day_{day}"):
                    selected_days.append(day)
        if st.checkbox("✅ تحديد الكل"):
            selected_days = days

        if selected_days and st.button(
            f"🚀 توليد {len(selected_days)} كراس", type="primary", use_container_width=True
        ):
            working_db = copy.deepcopy(db)
            generated = {}
            bar = st.progress(0)
            for idx, day in enumerate(selected_days):
                bar.progress(idx / len(selected_days))
                template = create_template(day)
                if template:
                    result, info, warns = build_daily_planner(
                        day, template, working_db, week_num, date_str
                    )
                    if result:
                        generated[day] = {'bytes': result, 'sessions': info, 'warnings': warns}
            bar.progress(1.0)
            st.session_state.gen = generated
            st.markdown(
                f'<div class="ok-box"><h3>✅ تم توليد {len(generated)} كراس!</h3></div>',
                unsafe_allow_html=True,
            )

        # التحميل والمعاينة
        gen_files = st.session_state.gen
        if gen_files:
            st.markdown("### 📥 تحميل الكراسات")
            download_cols = st.columns(min(len(gen_files), 5))
            for i, (day, data) in enumerate(gen_files.items()):
                with download_cols[i % 5]:
                    st.download_button(
                        f"📄 {day}",
                        data=data['bytes'],
                        file_name=f"كراس_{day}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"download_{day}",
                    )

            st.markdown("---")
            st.markdown("### 👁️ معاينة")
            preview_day = st.selectbox("اختر اليوم للمعاينة", list(gen_files.keys()))
            if preview_day:
                for session in gen_files[preview_day]['sessions']:
                    session_type = session['نوع']
                    css = {'روتيني': 'slot-r', 'تعليمي': 'slot-t', 'ناقص': 'slot-w'}.get(session_type, 'slot-r')
                    icon = {'روتيني': '⏰', 'تعليمي': '📖', 'ناقص': '⚠️'}.get(session_type, '⏰')
                    domain = session.get('المجال', '—')
                    badge = domain_badge(domain) if domain != '—' else ''
                    extra = ''
                    if session_type == 'تعليمي':
                        extra = (f"<br><small>📝 {session['الموضوع']}</small>"
                                 f"<br><small>🎯 {session['الكفاءة']}</small>")
                    st.markdown(
                        f'<div class="slot {css}">'
                        f'<span style="font-size:1.2rem">{icon}</span>'
                        f'<div style="flex:1"><strong>{session["النشاط"]}</strong> '
                        f'<span style="color:#888">({session["المدة"]})</span>'
                        f'{badge}{extra}</div></div>',
                        unsafe_allow_html=True,
                    )

# ──── تبويب 3: الإعداد ────
with tab3:
    st.markdown(f"### ⚙️ تخصيص مواد {current_level}")
    cfg = get_cfg()
    subjects_data = [{"المادة": n, "المجال": i["ج"], "الحصص": i["ح"]} for n, i in cfg["مو"].items()]
    edited_data = st.data_editor(
        subjects_data,
        num_rows="dynamic",
        column_config={
            "المادة": st.column_config.TextColumn("المادة", width="medium"),
            "المجال": st.column_config.SelectboxColumn("المجال", options=list(cfg["مج"].keys()), width="medium"),
            "الحصص": st.column_config.NumberColumn("الحصص", min_value=1, max_value=10, width="small"),
        },
        use_container_width=True,
        key="subjects_editor",
    )

    btn_col1, btn_col2 = st.columns(2)
    with btn_col1:
        if st.button("💾 حفظ التعديلات", use_container_width=True, type="primary"):
            new_subjects = {}
            for row in edited_data:
                if row.get("المادة") and row.get("المجال"):
                    new_subjects[row["المادة"]] = {"ج": row["المجال"], "ح": row.get("الحصص", 1)}
            if new_subjects:
                st.session_state.cfgs[current_level]["مو"] = new_subjects
                st.success("✅ تم حفظ التعديلات!")
    with btn_col2:
        if st.button("🔄 توليد توقيت تلقائي", use_container_width=True):
            new_cfg = get_cfg()
            new_schedule = auto_schedule(new_cfg["مو"], new_cfg["رت"])
            st.session_state.cfgs[current_level]["تو"] = new_schedule
            st.success("✅ تم توليد التوقيت!")

    total_edited = sum(r.get("الحصص", 0) for r in edited_data if r.get("المادة"))
    st.info(f"📊 المجموع: **{total_edited}** حصة/أسبوع (المتاح: ~34)")

# ──── تبويب 4: المجالات والتوقيت ────
with tab4:
    col_left, col_right = st.columns(2)

    with col_left:
        st.markdown("### 🗺️ المجالات التعليمية")
        subjects = get_subjects()
        domains_grouped = {}
        for subj, info in subjects.items():
            domains_grouped.setdefault(info["ج"], []).append(subj)
        for domain_name, domain_subjects in domains_grouped.items():
            color = domain_color(domain_name)
            total_hours = sum(subjects[s]["ح"] for s in domain_subjects)
            st.markdown(
                f'<div style="border:2px solid {color};border-radius:12px;padding:.8rem;margin:.4rem 0;">'
                f'<h4 style="color:{color};text-align:center;">{domain_name} ({total_hours} ح/أسبوع)</h4>',
                unsafe_allow_html=True,
            )
            for subj in domain_subjects:
                count = subjects[subj]["ح"]
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:6px;margin:3px 0;direction:rtl;">'
                    f'<span style="min-width:130px;font-size:.85rem">{subj}</span>'
                    f'<div style="background:{color}44;border-radius:3px;height:16px;width:{count * 14}px;"></div>'
                    f'<span style="color:{color};font-weight:700;font-size:.85rem">{count}</span></div>',
                    unsafe_allow_html=True,
                )
            st.markdown('</div>', unsafe_allow_html=True)

    with col_right:
        st.markdown("### 📊 التوقيت الأسبوعي")
        schedule = get_schedule()
        if schedule:
            view_day = st.selectbox("اختر اليوم", list(schedule.keys()), key="schedule_view")
            if view_day:
                routine = get_routine()
                for period_name, period_code in [("☀️ الفترة الصباحية", "ص"), ("🌙 الفترة المسائية", "م")]:
                    sessions = [s for s in schedule[view_day] if s["ف"] == period_code]
                    if not sessions:
                        continue
                    st.markdown(f"**{period_name}**")
                    table_data = [{
                        "#": j,
                        "النشاط": s['ن'],
                        "المدة": s['م'],
                        "المجال": domain_for(s['ن']) if s['ن'] not in routine else "—",
                    } for j, s in enumerate(sessions, 1)]
                    st.dataframe(table_data, use_container_width=True, hide_index=True)
        else:
            st.warning("💡 اضغط 'توليد توقيت تلقائي' في تبويب الإعداد")
