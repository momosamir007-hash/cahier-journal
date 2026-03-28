# -*- coding: utf-8 -*-
"""
إنشاء قالب الكراس اليومي template.docx
شغّل هذا الملف مرة واحدة فقط
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_rtl(cell):
    """ضبط اتجاه الخلية من اليمين لليسار"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = paragraph._p.get_or_add_pPr()
        bidi = pPr.makeelement(qn('w:bidi'), {})
        pPr.append(bidi)


def set_cell_text(cell, text, bold=False, size=10):
    """كتابة نص في خلية مع تنسيق"""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Sakkal Majalla"
    # ضبط الخط العربي
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:cs'): 'Sakkal Majalla'
    })
    rPr.append(rFonts)
    set_cell_rtl(cell)


def set_doc_rtl(doc):
    """ضبط اتجاه المستند بالكامل RTL"""
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = sectPr.makeelement(qn('w:bidi'), {})
        sectPr.append(bidi)


def shade_cells(row, color="D9E2F3"):
    """تلوين خلفية صف كامل"""
    for cell in row.cells:
        shading = cell._tc.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): color,
            qn('w:val'): 'clear'
        })
        shading.append(shd)


def create_period_table(doc, title, start_index, num_sessions):
    """إنشاء جدول لفترة واحدة (صباحية أو مسائية)"""
    # عنوان الفترة
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # إنشاء الجدول
    # الأعمدة: مؤشرات الكفاءة | عنوان الدرس | الميدان | النشاط | المدة
    table = doc.add_table(rows=1 + num_sessions, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── ضبط عرض الأعمدة ──
    col_widths = [Cm(5.5), Cm(5), Cm(3.5), Cm(3.5), Cm(2)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # ── صف العناوين ──
    headers = ['مؤشرات الكفاءة', 'عنوان الدرس', 'الميدان', 'النشاط', 'المدة']
    header_row = table.rows[0]
    shade_cells(header_row, "1F4E79")
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_rtl(cell)

    # ── صفوف الحصص مع المفاتيح ──
    for j in range(num_sessions):
        session_num = start_index + j
        data_row = table.rows[1 + j]
        # تلوين متناوب
        if j % 2 == 0:
            shade_cells(data_row, "EDF2F9")
        placeholders = [
            f'{{{{كفاءة_{session_num}}}}}',
            f'{{{{موضوع_{session_num}}}}}',
            f'{{{{ميدان_{session_num}}}}}',
            f'{{{{نشاط_{session_num}}}}}',
            f'{{{{مدة_{session_num}}}}}',
        ]
        for i, placeholder in enumerate(placeholders):
            set_cell_text(data_row.cells[i], placeholder, size=9)

    return table


def create_word_template():
    """الدالة الرئيسية لإنشاء القالب"""
    doc = Document()
    set_doc_rtl(doc)

    # ══════════════════════════════════════
    # ترويسة الكراس
    # ══════════════════════════════════════
    # الجمهورية
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run('الجمهورية الجزائرية الديمقراطية الشعبية')
    r1.bold = True
    r1.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('وزارة التربية الوطنية')
    r2.font.size = Pt(11)

    # عنوان الكراس
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run('الكراس اليومي')
    rt.bold = True
    rt.font.size = Pt(18)
    rt.font.color.rgb = RGBColor(0, 51, 102)

    # معلومات اليوم
    info_table = doc.add_table(rows=1, cols=3)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(
        info_table.rows[0].cells[2],
        'اليوم : {{اليوم}}',
        bold=True,
        size=12
    )
    set_cell_text(
        info_table.rows[0].cells[1],
        'التاريخ : {{التاريخ}}',
        size=11
    )
    set_cell_text(
        info_table.rows[0].cells[0],
        'الأسبوع : {{الأسبوع}}',
        size=11
    )
    doc.add_paragraph('')  # مسافة

    # ══════════════════════════════════════
    # جدول الفترة الصباحية (7 حصص)
    # ══════════════════════════════════════
    create_period_table(
        doc,
        title='☀ الفترة الصباحية',
        start_index=1,
        num_sessions=7
    )
    doc.add_paragraph('')  # فاصل

    # ══════════════════════════════════════
    # جدول الفترة المسائية (5 حصص)
    # ══════════════════════════════════════
    create_period_table(
        doc,
        title='🌙 الفترة المسائية',
        start_index=8,
        num_sessions=5
    )
    doc.add_paragraph('')  # مسافة

    # ══════════════════════════════════════
    # تذييل الملاحظات
    # ══════════════════════════════════════
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rn = note.add_run(
        'ملاحظات : '
        '................................................................'
        '................................................................'
    )
    rn.font.size = Pt(10)

    # ══════════════════════════════════════
    # الحفظ
    # ══════════════════════════════════════
    output = 'template.docx'
    doc.save(output)

    print(f"""
╔══════════════════════════════════════════════╗
║ ✅ تم إنشاء القالب بنجاح!                   ║
║ 📄 الملف: {output:<33}║
║                                             ║
║ المفاتيح المستخدمة:                        ║
║ صباح: {{{{نشاط_1}}}} → {{{{نشاط_7}}}}         ║
║ مساء: {{{{نشاط_8}}}} → {{{{نشاط_12}}}}        ║
║                                             ║
║ يمكنك فتحه بـ Word وتعديل الشكل             ║
║ مع الحفاظ على المفاتيح {{{{...}}}}              ║
╚══════════════════════════════════════════════╝
""")


if __name__ == "__main__":
    create_word_template()
